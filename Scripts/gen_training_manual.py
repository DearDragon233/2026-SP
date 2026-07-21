#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Generate: 育种模型项目技术培训手册.docx
Complete training manual for new coworkers
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
style.font.size = Pt(12)
style.font.name = 'Times New Roman'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, indent=False, font_size=12):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_code(text):
    """Add code block with monospace font"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_table(headers, rows, col_widths=None):
    """Add a formatted table"""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'E8E8E8')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
    
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                if ci < len(row.cells):
                    row.cells[ci].width = Cm(w)
    
    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('环境-作物-管理互作模型')
run.bold = True
run.font.size = Pt(26)
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_p.add_run('技术培训手册：从入门到协作开发')
run.font.size = Pt(18)
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(4):
    doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_p.add_run('彭宇程 | 2026年7月 | 版本 1.0')
run.font.size = Pt(14)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TOC placeholder
# ═══════════════════════════════════════════════════════════════
add_heading_styled('目录', 1)
add_para('本手册面向初次加入项目的coworker，覆盖从环境搭建到模型部署的完整技术链路。每个模块独立成章，可按需阅读，也可从头到尾系统学习。', indent=True)

toc_items = [
    ('第一部分：项目总览与协作基础', [
        '第1章  项目架构与分工设计',
        '第2章  Git与GitHub团队协作',
        '第3章  Python开发环境与项目管理',
    ]),
    ('第二部分：数据获取与预处理', [
        '第4章  地理空间数据处理（rasterio/xarray/geopandas）',
        '第5章  多源数据融合与特征工程（pandas/numpy）',
        '第6章  环境指纹构建',
    ]),
    ('第三部分：建模与优化', [
        '第7章  机器学习建模：从决策树到XGBoost',
        '第8章  模型解释：SHAP值原理与应用',
        '第9章  多目标优化：NSGA-II',
        '第10章 不确定性量化：Bootstrap方法',
    ]),
    ('第四部分：R语言统计与可视化', [
        '第11章 R语言统计分析',
        '第12章 R语言科学可视化',
    ]),
    ('第五部分：部署与可重复性', [
        '第13章 模型部署与API服务（FastAPI）',
        '第14章 容器化与可重复研究（Docker）',
        '第15章 学术写作与文献管理',
    ]),
    ('附录', [
        '附录A  推荐学习路径（3周速成方案）',
        '附录B  常见问题排查',
        '附录C  术语对照表',
    ]),
]

for part, chapters in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(part)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    for ch in chapters:
        add_para(f'    {ch}', font_size=11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# PART 1: PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════

# ── Chapter 1 ──
add_heading_styled('第一部分  项目总览与协作基础', 1)
add_heading_styled('第1章  项目架构与分工设计', 2)

add_heading_styled('1.1  你能学到什么', 3)
add_para('本章介绍整个项目的技术架构、模块划分和团队分工方式。读完本章，你将理解：一个多模块数据科学项目如何从零组织；各模块之间的数据流和依赖关系；如何在团队中分配任务而不产生耦合冲突。', indent=True)

add_heading_styled('1.2  项目背景', 3)
add_para('本项目旨在利用公开数据，构建一个"环境指纹—品种性状—管理方案"三要素互作模型。对标 Xiao et al. (2024, Nature Food 5:59-71) 的研究框架，我们使用世界气候数据(WorldClim)、土壤数据(SoilGrids, 中国土壤数据库)、地形数据(SRTM)、大气CO2浓度(CMIP6)和作物品种审定公告，构建约65维输入特征矩阵，通过XGBoost模拟APSIM过程模型输出，再使用NSGA-II搜索最优管理方案，并用SHAP值解释各特征对产量和环境指标的边际贡献。', indent=True)
add_para('项目的最终产出包括两套输出：最优农艺管理推荐方案（氮肥、灌溉、秸秆还田）；最优育种改良方向（哪些性状最值得改良，改良幅度与产量增益的关系）。', indent=True)

add_heading_styled('1.3  总体技术架构', 3)
add_para('整个项目按照数据流向分为五个层次：', indent=True)
add_para('第一层：数据获取层。负责从公开数据源下载原始栅格/表格数据，存放于D盘2026-SP/Data目录下。涉及Python的rasterio、urllib、BeautifulSoup等库和手动浏览器下载。', indent=True)
add_para('第二层：特征工程层。将不同分辨率、不同投影的多源数据统一到1km网格，构建标准化的环境指纹矩阵。核心工具为rasterio（重采样）、geopandas（空间连接）、pandas（表格操作）。', indent=True)
add_para('第三层：建模层。使用XGBoost回归器训练从环境指纹到产量/环境指标的映射模型。核心工具为xgboost、scikit-learn（交叉验证、超参数调优）。', indent=True)
add_para('第四层：优化层。将训练好的XGBoost模型作为代理模型（surrogate），使用NSGA-II多目标遗传算法搜索最优的管理方案组合。核心工具为pymoo。', indent=True)
add_para('第五层：解释与部署层。使用SHAP解释特征重要性，使用Bootstrap量化不确定性，使用FastAPI将模型部署为RESTful服务。', indent=True)

add_heading_styled('1.4  团队分工设计', 3)
add_para('建议按模块分为四个角色，每个角色可由1-2人承担：', indent=True)

add_table(
    ['角色', '负责模块', '核心技能', '交付物'],
    [
        ['数据工程师', '第1-2层：数据获取、预处理、特征工程', 'rasterio, geopandas, pandas, xarray', '环境指纹矩阵(CSV/Parquet)'],
        ['建模工程师', '第3-4层：XGBoost训练、NSGA-II优化', 'xgboost, scikit-learn, pymoo, optuna', '训练好的模型(pkl)、优化结果CSV'],
        ['分析工程师', '第5层：SHAP解释、Bootstrap、统计分析', 'shap, numpy, scipy, R/tidyverse', '特征重要性图、不确定性区间、统计检验报告'],
        ['部署工程师', '第5层：API服务、可视化、文档', 'FastAPI, Docker, GitHub Actions', 'Docker镜像、API文档、使用手册'],
    ],
    [4, 5, 5, 5]
)

add_heading_styled('1.5  模块间接口约定', 3)
add_para('为避免模块间的耦合冲突，我们约定以下数据接口格式：所有模块间传递的数据均为CSV或Parquet格式的表格文件；每一行代表一个地理位置点（网格）；列名使用英文小写+下划线命名（snake_case）；缺失值统一用NaN表示，不允许使用-999等魔法数字；空间参考统一为WGS84(EPSG:4326)，分辨率统一为0.008333度（约1km）。', indent=True)

add_heading_styled('1.6  掌握后的延展能力', 3)
add_para('掌握了本章所述的项目架构设计方法后，你可以类推到任何以数据为中心的科研项目组织：多源遥感数据分析、环境监测系统、作物模型区域升尺度、以及任何需要"数据获取—特征工程—建模—优化—部署"五步链路的应用场景。这种分层架构是工业界数据科学团队的标准实践，掌握它就具备了在科技公司担任数据工程师或机器学习工程师的基础架构能力。', indent=True)
add_para('当前业界趋势（2025-2026）：模块化ML管道（Modular ML Pipelines）正成为主流。Kubeflow、MLflow、Metaflow等工具本质上都是对这种分层架构的工程化实现。你在本项目中学到的"约定接口+独立模块"思想，与这些工业级工具的设计哲学完全一致。', indent=True)

doc.add_page_break()

# ── Chapter 2 ──
add_heading_styled('第2章  Git与GitHub团队协作', 2)

add_heading_styled('2.1  你能学到什么', 3)
add_para('本章介绍使用Git和GitHub进行团队代码管理的方法。读完本章，你将能够：创建和克隆代码仓库；使用分支(branch)进行独立开发而不影响他人；提交合并请求(Pull Request)并接受代码审查；解决合并冲突(merge conflict)；使用GitHub Projects或Issues跟踪任务进度。', indent=True)

add_heading_styled('2.2  为什么Git是必备技能', 3)
add_para('Git是目前全球使用率最高的版本控制系统，2024年Stack Overflow开发者调查显示超过93%的专业开发者使用Git。在科研领域，Nature和Science均建议论文代码通过GitHub发布以保证可重复性。对个人而言，GitHub profile已成为科技行业求职的事实标准简历；对团队而言，Git分支模型是多人协作而不产生代码冲突的唯一可靠方案。', indent=True)

add_heading_styled('2.3  项目仓库结构', 3)
add_para('推荐在GitHub上创建Organization或公共仓库，结构如下：', indent=True)
add_code('2026-SP/')
add_code('├── .github/            # GitHub Actions CI/CD 配置')
add_code('│   └── workflows/')
add_code('├── data/              # 数据目录（gitignore，不上传大文件）')
add_code('├── src/               # 源代码')
add_code('│   ├── data_ingest/   # 数据获取脚本')
add_code('│   ├── features/      # 特征工程')
add_code('│   ├── models/        # 建模与训练')
add_code('│   ├── optimization/  # NSGA-II优化')
add_code('│   └── visualization/ # 可视化')
add_code('├── tests/             # 单元测试')
add_code('├── docs/              # 文档')
add_code('├── notebooks/         # Jupyter探索性分析')
add_code('├── requirements.txt   # Python依赖')
add_code('├── Dockerfile         # 容器化配置')
add_code('├── README.md          # 项目说明')
add_code('└── .gitignore         # 忽略文件列表')
p = doc.add_paragraph()
run = p.add_run('注：大文件（>100MB）如GeoTIFF不应提交到Git，应使用.gitignore排除，通过README中的下载链接获取，或使用Git LFS（Large File Storage）。')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.color.rgb = RGBColor(100, 100, 100)

add_heading_styled('2.4  分支策略', 3)
add_para('推荐GitHub Flow模型，这是目前最简洁且工业界最广泛使用的分支策略：', indent=True)
add_para('main分支：永远保持可运行状态，只通过PR合并，禁止直接push。每次合并到main应触发自动测试。', indent=True)
add_para('feature分支：从main分支创建，命名为feature/功能描述，例如feature/add-nsga2-optimizer、feature/shap-visualization。一个分支只做一件事，完成后发起PR合并回main。', indent=True)
add_para('fix分支：紧急修复分支，命名为fix/问题描述，例如fix/memory-leak-in-rasterio。修复后同时合并到main。', indent=True)

add_heading_styled('2.5  日常工作流', 3)
add_para('每天开始工作时的标准操作：', indent=True)
add_code('git checkout main           # 切换到main分支')
add_code('git pull origin main       # 拉取最新代码')
add_code('git checkout -b feature/xxx  # 创建新功能分支')
add_para('完成一个功能后的提交流程：', indent=True)
add_code('git add .                   # 暂存所有修改')
add_code('git commit -m "feat: 添加NSGA-II优化模块"  # 提交')
add_code('git push origin feature/xxx # 推送到远程')
add_para('然后在GitHub网页上创建Pull Request，指定至少一位reviewer。审查通过后点击Merge，最后删除远程分支：', indent=True)
add_code('git branch -d feature/xxx   # 删除本地分支')
add_code('git push origin --delete feature/xxx  # 删除远程分支')

add_heading_styled('2.6  提交信息规范', 3)
add_para('统一使用Conventional Commits格式，让提交历史清晰可读：', indent=True)
add_table(
    ['前缀', '含义', '示例'],
    [
        ['feat', '新功能', 'feat: 添加SHAP特征重要性分析模块'],
        ['fix', '修复bug', 'fix: 修复rasterio内存泄漏问题'],
        ['docs', '文档更新', 'docs: 更新API接口文档'],
        ['refactor', '代码重构', 'refactor: 将数据加载逻辑提取为独立模块'],
        ['test', '测试相关', 'test: 添加XGBoost交叉验证单元测试'],
        ['chore', '杂项/构建', 'chore: 更新requirements.txt依赖版本'],
    ],
    [2.5, 3, 11]
)

add_heading_styled('2.7  掌握后的延展能力', 3)
add_para('掌握Git/GitHub工作流后，你可以参与任何开源项目（从提交第一个PR开始），理解CI/CD持续集成的基本概念，为后续学习DevOps打基础。当前趋势（2025-2026）是"GitOps"——将Git作为基础设施配置和部署的唯一真相来源(source of truth)，Kubernetes生态的ArgoCD和Flux CD都基于这一理念。你在这里学到的分支和PR流程，直接复用于这些前沿工具。', indent=True)

doc.add_page_break()

# ── Chapter 3 ──
add_heading_styled('第3章  Python开发环境与项目管理', 2)

add_heading_styled('3.1  你能学到什么', 3)
add_para('本章介绍如何搭建一个可复现的Python开发环境。你将学会使用Conda管理虚拟环境、使用pip安装依赖、理解requirements.txt的结构、配置VS Code进行高效开发。这些都是所有后续模块的基础。', indent=True)

add_heading_styled('3.2  环境隔离的重要性', 3)
add_para('Python生态最大的痛点之一是依赖冲突：项目A需要numpy 1.24，项目B需要numpy 2.0，直接装在系统Python中必然导致其中一个项目报错。虚拟环境（virtual environment）为每个项目创建独立的Python解释器和包安装目录，彻底解决这一问题。', indent=True)

add_heading_styled('3.3  推荐环境方案', 3)
add_para('方案一：Miniconda + conda环境（推荐新手使用）。安装Miniconda（轻量版Anaconda，约50MB），然后：', indent=True)
add_code('conda create -n 2026sp python=3.11   # 创建环境')
add_code('conda activate 2026sp              # 激活环境')
add_code('pip install numpy pandas xarray rasterio geopandas xgboost shap scikit-learn pymoo optuna matplotlib seaborn jupyter fastapi uvicorn docker')
add_code('pip freeze > requirements.txt      # 锁定依赖版本')
add_para('方案二：Python venv + pip（轻量方案）。适合已有Python 3.11+的用户：', indent=True)
add_code('python -m venv venv                # 创建虚拟环境')
add_code('venv\\Scripts\\activate             # Windows激活')
add_code('pip install -r requirements.txt    # 安装所有依赖')

add_heading_styled('3.4  项目文件组织规范', 3)
add_para('每个功能模块的Python文件应遵循以下规范：文件头包含模块说明docstring；使用if __name__ == "__main__"保护可执行代码；将可配置参数抽取到config.py或config.yaml中；函数和变量命名使用snake_case；类命名使用PascalCase。', indent=True)
add_para('示例：一个规范的特征工程模块开头：', indent=True)
add_code('"""')
add_code('特征工程模块：从原始栅格数据构建环境指纹矩阵')
add_code('')
add_code('输入：Data/WorldClim/*.tif, Data/SRTM/*.tif, ...')
add_code('输出：features_matrix.parquet')
add_code('用法：python -m src.features.build_matrix --config config.yaml')
add_code('"""')
add_code('import rasterio')
add_code('import numpy as np')
add_code('import pandas as pd')
add_code('from pathlib import Path')
add_code('')
add_code('')
add_code('def extract_features(grid_points, raster_dir):')
add_code('    """从栅格数据中提取每个网格点的环境特征"""')
add_code('    pass')
add_code('')
add_code('')
add_code('if __name__ == "__main__":')
add_code('    main()')

add_heading_styled('3.5  掌握后的延展能力', 3)
add_para('虚拟环境和依赖管理是所有Python项目的起点。掌握了这些，你可以无缝切换到任何Python数据科学项目，包括深度学习（PyTorch/TensorFlow环境）、Web开发（Django/Flask环境）、以及数据工程（Apache Spark/PySpark环境）。当前趋势（2025-2026）是使用uv或pixi等新一代Python包管理器，它们比pip快10-100倍，但底层原理完全相同——你在这里学到的概念可以直接迁移。', indent=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# PART 2: DATA
# ═══════════════════════════════════════════════════════════════
add_heading_styled('第二部分  数据获取与预处理', 1)

# ── Chapter 4 ──
add_heading_styled('第4章  地理空间数据处理', 2)
add_para('核心库：rasterio, xarray, geopandas, rioxarray', indent=True)

add_heading_styled('4.1  你能学到什么', 3)
add_para('本章是项目的数据基础。你将学会：用rasterio读取和写入GeoTIFF格式的栅格数据；用xarray处理NetCDF格式的多维气候数据；用geopandas处理矢量边界（如平谷区行政边界）并进行空间连接；将不同分辨率、不同投影的数据统一重采样到同一网格。', indent=True)

add_heading_styled('4.2  栅格数据的本质', 3)
add_para('GeoTIFF文件本质上是一个二维（或三维）数字矩阵，附带了地理参考信息。地理参考信息包括：投影坐标系（CRS，Coordinate Reference System），定义了如何将经纬度映射到平面；仿射变换矩阵（Affine Transform），定义了像素坐标(row, col)与地理坐标(lon, lat)之间的线性转换关系；以及NoData值，标记无效区域。', indent=True)
add_para('当你用rasterio.open()打开一个GeoTIFF文件时，你获得的是一个类似NumPy数组的对象，但它"知道"自己在地球上的位置。这是地理空间数据分析的基本抽象。', indent=True)

add_heading_styled('4.3  核心操作', 3)
add_para('操作一：读取栅格数据并获取元信息', indent=True)
add_code('import rasterio')
add_code('with rasterio.open("wc2.1_2.5m_bio_1.tif") as src:')
add_code('    print(f"CRS: {src.crs}")')
add_code('    print(f"分辨率: {src.res}")')
add_code('    print(f"尺寸: {src.width} x {src.height}")')
add_code('    print(f"边界: {src.bounds}")')
add_code('    data = src.read(1)  # 读取第一波段')
add_para('操作二：按经纬度提取像素值。给定目标点的经纬度，将经纬度转换为像素行列号，再读取该位置的数值：', indent=True)
add_code('from rasterio.transform import rowcol')
add_code('row, col = rowcol(src.transform, lon, lat)')
add_code('value = data[row, col]  # 该位置的年均温')
add_para('操作三：重采样到统一分辨率。不同数据源的分辨率不同（WorldClim约5km，SRTM约30m），需要统一重采样到1km网格：', indent=True)
add_code('from rasterio.enums import Resampling')
add_code('# 计算1km分辨率下的新尺寸')
add_code('new_height = int(src.height * src.res[1] / 0.008333)')
add_code('new_width = int(src.width * src.res[0] / 0.008333)')
add_code('resampled = src.read(1, out_shape=(new_height, new_width),')
add_code('                        resampling=Resampling.bilinear)')

add_heading_styled('4.4  与项目管理衔接', 3)
add_para('地理空间数据处理模块的输出是"每个目标网格点的特征值向量"。这个输出将被第5章的特征工程模块消费。因此，模块的输出格式约定为：一个Pandas DataFrame，列包含grid_id、lon、lat、bio1、bio2 ... 等；存储为Parquet格式（比CSV快5-10倍，且保留数据类型）；每个数据源单独一个文件，最后在第5章中合并。', indent=True)
add_para('在GitHub仓库中，数据处理脚本放在src/data_ingest/目录下。每个公开数据源对应一个子模块，例如：src/data_ingest/worldclim.py、src/data_ingest/soilgrids.py、src/data_ingest/srtm.py。', indent=True)

add_heading_styled('4.5  常见问题与优化', 3)
add_para('内存溢出：读取全球范围的GeoTIFF可能导致内存不足。解决方案是使用窗口读取(windowed reading)，只加载目标区域：', indent=True)
add_code('from rasterio.windows import from_bounds')
add_code('window = from_bounds(116.7, 40.0, 117.5, 40.5, src.transform)')
add_code('subset = src.read(1, window=window)')
add_para('投影不一致：不同数据源的CRS可能不同（WGS84 vs Goode Homolosine vs  Albers），必须统一转换。rasterio.warp.reproject可以将一个栅格重新投影到目标CRS。', indent=True)
add_para('处理速度优化：对于大规模批量提取，不要逐个点循环读取，而应一次性读取整个目标区域的numpy数组，然后用向量化索引提取所有点。循环读取（10000次I/O操作）可能耗时数分钟，而一次性读取+向量化索引只需几秒。', indent=True)

add_heading_styled('4.6  掌握后的延展能力', 3)
add_para('地理空间数据处理是遥感、GIS、环境科学和精准农业的核心技能。掌握了rasterio/geopandas/xarray这套工具链后，你可以处理任何卫星遥感数据（Landsat、Sentinel-2、MODIS）、气候再分析数据（ERA5、NCEP）、以及数字土壤制图产品。当前趋势（2025-2026）是"云原生地理空间"（Cloud-Native Geospatial）：STAC（SpatioTemporal Asset Catalog）规范和COG（Cloud Optimized GeoTIFF）格式使得无需下载完整文件即可在云端查询和读取特定时空范围的数据。你在这里学到的rasterio窗口读取技术，正是这一趋势的基础。', indent=True)

doc.add_page_break()

# ── Chapter 5 ──
add_heading_styled('第5章  多源数据融合与特征工程', 2)
add_para('核心库：pandas, numpy, scipy', indent=True)

add_heading_styled('5.1  你能学到什么', 3)
add_para('本章教你如何将多个不同来源、不同格式的数据融合为统一的分析就绪表格。你将掌握：Pandas DataFrame的高级操作（merge/join/groupby/transform）；处理缺失值的多种策略（删除、均值填充、KNN填充、MICE多重插补）；特征缩放与标准化（StandardScaler、MinMaxScaler）；衍生特征的计算（从已有特征生成新特征）。', indent=True)

add_heading_styled('5.2  为什么特征工程决定模型上限', 3)
add_para('数据科学界有一句名言："数据和特征决定了机器学习的上限，而模型和算法只是在逼近这个上限。"在环境-作物建模中尤其如此：如果土壤有机碳数据缺失或错误，无论用多复杂的模型都无法正确预测产量。特征工程是连接"原始数据"和"可用模型输入"之间的桥梁，也是项目中最耗时但最重要的环节（通常占整个项目时间的60-70%）。', indent=True)

add_heading_styled('5.3  数据融合流程', 3)
add_para('本项目的多源数据融合分为四步：', indent=True)
add_para('第一步：统一空间网格。以平谷区行政边界为掩膜，生成间距约1km的规则网格点（约65个点），每个点用(lon, lat)标识。所有后续提取均基于这个网格。', indent=True)
add_para('第二步：逐源提取特征。对每个网格点，从各个栅格数据源中提取对应的数值：从WorldClim bio tif中提取19个生物气候变量；从WorldClim monthly tif中提取3-9月的月均温和月降水（14维）；从SRTM tif中提取海拔、坡度和坡向（3维）；从SoilGrids VRT中提取砂粒、粉粒、黏粒、有机碳等（5维已获取+3维待补充）；从CO2查表中按年份匹配浓度（1维）。', indent=True)
add_para('第三步：衍生特征计算。从月温数据中计算生长度日（GDD, Growing Degree Days）：对每个基温阈值（0度C、5度C、10度C、15度C、20度C、25度C），模拟逐日温度曲线（三角函数插值）并累加超过阈值的热量单位。GDD是作物模型中最重要的热时间指标。另从CHELSA逐日数据中计算极端气候指标：生育期Tmax的95分位数（热胁迫）、Tmin的5分位数（冷胁迫）、降水的95分位数（涝渍风险）。', indent=True)
add_para('第四步：合并与清洗。将所有提取结果按(grid_id, year)合并为一张大表，处理缺失值，检查异常值（如海拔为负在平谷区不合理），进行标准化。', indent=True)

add_heading_styled('5.4  核心代码示例', 3)
add_para('多源数据合并（merge）的核心操作：', indent=True)
add_code('import pandas as pd')
add_code('')
add_code('# 加载各数据源（均按grid_id索引）')
add_code('df_bio = pd.read_parquet("features_bio.parquet")')
add_code('df_clim = pd.read_parquet("features_clim.parquet")')
add_code('df_soil = pd.read_parquet("features_soil.parquet")')
add_code('df_terrain = pd.read_parquet("features_terrain.parquet")')
add_code('')
add_code('# 链式合并')
add_code('df_all = (df_bio')
add_code('    .merge(df_clim, on=["grid_id", "year"])')
add_code('    .merge(df_soil, on="grid_id")')
add_code('    .merge(df_terrain, on="grid_id")')
add_code(')')
add_para('GDD计算（从月均温通过三角函数插值得到逐日温度，再累加超过基温的部分）：', indent=True)
add_code('import numpy as np')
add_code('')
add_code('def calc_gdd(monthly_tavg, base_temp):')
add_code('    """从12个月的月均温计算生长度日"""')
add_code('    # 三角函数重建逐日温度（Forsythe et al. 1995方法）')
add_code('    days_per_month = [31,28,31,30,31,30,31,31,30,31,30,31]')
add_code('    daily_temps = []')
add_code('    for m in range(12):')
add_code('        # 月内每日温度正弦插值')
add_code('        for d in range(days_per_month[m]):')
add_code('            frac = (d + 0.5) / days_per_month[m]')
add_code('            daily_temps.append(monthly_tavg[m])')
add_code('    daily_temps = np.array(daily_temps)')
add_code('    return np.maximum(daily_temps - base_temp, 0).sum()')
add_para('缺失值处理的三层策略：', indent=True)
add_code('from sklearn.impute import SimpleImputer, KNNImputer')
add_code('')
add_code('# 策略1: 若缺失率 < 1%, 中位数填充')
add_code('imputer = SimpleImputer(strategy="median")')
add_code('')
add_code('# 策略2: 若缺失率 1-10%, KNN填充（利用相邻网格相关性的填充）')
add_code('imputer = KNNImputer(n_neighbors=5)')
add_code('')
add_code('# 策略3: 若缺失率 > 10%, 该特征考虑剔除或在论文中明确讨论')

add_heading_styled('5.5  掌握后的延展能力', 3)
add_para('特征工程是数据科学最通用的技能，几乎适用于任何领域。你在这里学到的pandas链式操作（merge-transform-groupby流水线）、缺失值处理策略、衍生特征构建方法，在以下领域完全复用：金融风控（交易特征构建）、推荐系统（用户行为特征）、医疗诊断（电子病历特征）、NLP（文本特征提取）。当前趋势（2025-2026）是"特征存储"（Feature Store）的兴起——像Feast和Tecton这样的工具将特征工程标准化为可共享、可版本化的资产。你在这里学到的"特征矩阵"概念，正是Feature Store的核心数据模型。', indent=True)

doc.add_page_break()

# ── Chapter 6 ──
add_heading_styled('第6章  环境指纹构建', 2)
add_para('核心库：rasterio, numpy, pandas, scipy', indent=True)

add_heading_styled('6.1  你能学到什么', 3)
add_para('本章将第4章和第5章的技能串联起来，完成项目中最核心的数据产物："环境指纹矩阵"。你将学会：如何将20余个独立的GeoTIFF文件转化为一张规整的分析表；如何设计一个可复用的特征提取流水线；如何验证提取结果的空间一致性和物理合理性。', indent=True)

add_heading_styled('6.2  什么是"环境指纹"', 3)
add_para('"环境指纹"（Environmental Fingerprint）是我们在本项目中的核心概念。它为每一个地理位置创建一个独特的"身份标识"——由55-75个数值组成的向量，完整地刻画该位置的气候、土壤、地形和大气条件。可以类比于：人类的指纹由几十个特征点唯一标识一个人；农田的"指纹"由几十个环境变量唯一标识一块地。', indent=True)
add_para('这个概念对应Xiao et al. (2024)论文中的"环境特征空间"。在论文中，作者用约30-35个环境变量训练XGBoost模型来模拟APSIM的输出。我们的"指纹"在此基础上扩展了品种性状维度和更精细的土壤属性，力图提供更全面的环境刻画。', indent=True)

add_heading_styled('6.3  结构程序设计', 3)
add_para('特征提取流水线应设计为模块化的Pipeline类：', indent=True)
add_code('class EnvironmentalFingerprintBuilder:')
add_code('    """环境指纹构建器：从栅格数据中提取每个网格点的特征向量"""')
add_code('')
add_code('    def __init__(self, config):')
add_code('        self.data_dir = Path(config["data_dir"])')
add_code('        self.grid_points = self._load_grid(config["grid_file"])')
add_code('        self.target_crs = "EPSG:4326"')
add_code('')
add_code('    def _load_grid(self, grid_file):')
add_code('        """加载目标网格点（平谷区1km网格）"""')
add_code('        return pd.read_parquet(grid_file)')
add_code('')
add_code('    def extract_worldclim_bio(self):')
add_code('        """提取19个生物气候变量"""')
add_code('        bio_dir = self.data_dir / "WorldClim"')
add_code('        features = {}')
add_code('        for i in range(1, 20):')
add_code('            tif_path = bio_dir / f"wc2.1_2.5m_bio_{i}.tif"')
add_code('            with rasterio.open(tif_path) as src:')
add_code('                for idx, row in self.grid_points.iterrows():')
add_code('                    features.setdefault(f"bio{i}", []).append(')
add_code('                        src.sample([(row["lon"], row["lat"])]).item()')
add_code('        return pd.DataFrame(features)')
add_code('')
add_code('    def extract_all(self):')
add_code('        """执行完整提取流水线"""')
add_code('        df_bio = self.extract_worldclim_bio()')
add_code('        df_clim = self.extract_monthly_climate()')
add_code('        df_terrain = self.extract_terrain()')
add_code('        df_soil = self.extract_soil()')
add_code('        df_co2 = self.extract_co2()')
add_code('        return pd.concat([df_bio, df_clim, df_terrain, df_soil, df_co2], axis=1)')

add_heading_styled('6.4  质量控制', 3)
add_para('每条特征提取完成后，必须进行以下验证：', indent=True)
add_para('数值范围检查：检查提取值是否在物理合理范围内（例如平谷区年均温应在5-15度C范围内，年降水应在300-800mm范围内）。超出范围的值可能指示坐标错误或数据源问题。', indent=True)
add_para('空间一致性检查：相邻网格点的特征值应该是渐变的。如果某个网格点的bio1（年均温）比相邻点高5度C以上，可能有异常。可以计算变异函数(variogram)或简单的相邻点差值来检测。', indent=True)
add_para('缺失网格检查：确保每个目标网格点都有对应的特征值。如果某些点在数据源的范围之外（如边界网格），应该在输出中明确标注。', indent=True)

add_heading_styled('6.5  掌握后的延展能力', 3)
add_para('"环境指纹"这个概念可以直接推广到以下场景：任何需要"将地理位置转化为固定维度特征向量"的空间预测任务（如物种分布建模、土壤碳储量预测、城市热岛分析）；精准农业中的管理区划（Management Zone Delineation）——将农田按环境指纹聚类，对不同的管理区采用不同方案。当前趋势（2025-2026）是"数字孪生"（Digital Twin）在农业中的应用：为每一个田块创建一个数字副本，持续更新其状态。环境指纹是数字孪生的静态部分，搭配时序遥感数据后即构成完整的数字孪生数据层。', indent=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# PART 3: MODELING
# ═══════════════════════════════════════════════════════════════
add_heading_styled('第三部分  建模与优化', 1)

# ── Chapter 7 ──
add_heading_styled('第7章  机器学习建模：从决策树到XGBoost', 2)
add_para('核心库：scikit-learn, xgboost, optuna', indent=True)

add_heading_styled('7.1  你能学到什么', 3)
add_para('本章是项目的核心建模部分。你将学会：理解决策树的分裂原理（信息增益、基尼系数）；理解集成学习为什么优于单棵树（Bagging与Boosting的区别）；掌握XGBoost的梯度提升机制及其关键超参数；使用交叉验证评估模型性能；使用Optuna进行自动超参数调优。', indent=True)

add_heading_styled('7.2  从决策树到梯度提升：完整推导', 3)
add_para('第一步：决策树的直觉。假设你要预测某块地的玉米产量，你手上有年均温、年降水、土壤有机碳三个特征。决策树的思路是不断地问"是/否"问题来将数据分组：年均温是否大于12度C？如果是，分到左枝；年降水是否大于500mm？如果是，再分到左枝...每次分裂选择那个能让两组产量差异最大化的特征和阈值。在数学上，这个"差异"用信息增益或基尼系数衡量——分裂后子节点的"纯度"减去分裂前的"纯度"。', indent=True)
add_para('第二步：单棵树的致命缺陷。单棵决策树极易过拟合：它可以把训练数据分到每个叶子只有一个样本，在训练集上完美但在测试集上一塌糊涂。想象你背下了一张考试卷的所有答案——你在这张卷子上满分，但换一套卷子就全错了。这就是过拟合。', indent=True)
add_para('第三步：集成学习的思想。与其依赖一棵"独裁者"树，不如让一群树民主投票。Bagging（Bootstrap Aggregating，如随机森林）的思路是：并行训练100棵树，每棵用不同的数据子集和特征子集，最后取所有树的平均预测值。这能大幅降低方差。', indent=True)
add_para('第四步：Boosting的革命。Bagging的100棵树是独立训练的，每棵树不知道其他树在做什么。Boosting（提升法，如XGBoost）改变了这个范式：树是串行训练的，每一棵新树的使命是修正前面所有树的残差（误差）。想象你在学射箭：第一箭偏左了5厘米——第二箭你就专门瞄准"比目标偏右5厘米"的位置——第三箭再修正第二箭的偏差...最终，所有箭的"合力"会非常接近靶心。这就是梯度提升的直觉：每一棵新树拟合的是前序模型预测值与真实值之间的负梯度（即残差）。', indent=True)
add_para('第五步：XGBoost的工程优化。传统的梯度提升（如sklearn的GradientBoostingRegressor）每次分裂需要遍历所有特征的所有可能阈值，计算量巨大。XGBoost引入了三项关键优化：加权分位数草图（Weighted Quantile Sketch），将连续特征离散化为分位数桶，大幅减少候选分裂点；缓存感知访问（Cache-aware Access），利用CPU缓存结构加速数据读取；正则化项引入叶节点权重和树深度惩罚，天然防过拟合。这些优化使得XGBoost比传统的GBRT快10倍以上，同时精度更高。', indent=True)

add_heading_styled('7.3  本项目中的建模设计', 3)
add_para('目标变量：本项目需要训练多个XGBoost模型。小麦产量模型，输入环境指纹+品种性状+管理参数，输出小麦产量；玉米产量模型，同理；氮淋溶模型，输出NO3-N leaching量；土壤有机碳变化模型，输出delta-SOC。', indent=True)
add_para('训练-验证-测试划分策略：由于空间自相关的存在（相邻网格的环境非常相似），不能简单地随机划分。必须使用空间分块交叉验证（Spatial Block CV）：将平谷区的65个网格按地理位置分为5块，每次用4块训练、1块验证，旋转5次。这样可以检测模型能否泛化到未见过的地理位置，而不是靠空间邻近性作弊。', indent=True)

add_heading_styled('7.4  核心代码', 3)
add_code('import xgboost as xgb')
add_code('from sklearn.model_selection import cross_val_score')
add_code('import optuna')
add_code('')
add_code('# XGBoost基础训练')
add_code('model = xgb.XGBRegressor(')
add_code('    n_estimators=500,        # 树的数量')
add_code('    max_depth=6,             # 每棵树的最大深度')
add_code('    learning_rate=0.05,      # 学习率（每一棵树的"步伐"大小）')
add_code('    subsample=0.8,           # 每棵树随机使用80%训练样本')
add_code('    colsample_bytree=0.8,    # 每棵树随机使用80%特征')
add_code('    reg_alpha=0.1,           # L1正则化（稀疏性）')
add_code('    reg_lambda=1.0,          # L2正则化（权重衰减）')
add_code('    random_state=42')
add_code(')')
add_code('model.fit(X_train, y_train)')
add_code('score = model.score(X_test, y_test)  # R-squared')
add_code('')
add_code('# Optuna自动超参数调优')
add_code('def objective(trial):')
add_code('    params = {')
add_code('        "max_depth": trial.suggest_int("max_depth", 3, 10),')
add_code('        "learning_rate": trial.suggest_float("lr", 0.01, 0.3),')
add_code('        "subsample": trial.suggest_float("subsample", 0.5, 1.0),')
add_code('        "colsample_bytree": trial.suggest_float("colsample", 0.5, 1.0),')
add_code('        "reg_alpha": trial.suggest_float("alpha", 1e-8, 1.0),')
add_code('        "reg_lambda": trial.suggest_float("lambda", 1e-8, 1.0),')
add_code('    }')
add_code('    model = xgb.XGBRegressor(**params, n_estimators=500,')
add_code('                              early_stopping_rounds=20)')
add_code('    scores = cross_val_score(model, X_train, y_train, cv=5,')
add_code('                             scoring="r2")')
add_code('    return scores.mean()')
add_code('')
add_code('study = optuna.create_study(direction="maximize")')
add_code('study.optimize(objective, n_trials=50)')
add_code('print(f"最佳超参数: {study.best_params}")')

add_heading_styled('7.5  训练技巧与常见陷阱', 3)
add_para('早停（Early Stopping）：训练过程中如果在验证集上连续20轮没有提升，自动停止。这能防止过拟合并节省训练时间。', indent=True)
add_para('特征重要性偏差：XGBoost默认的feature_importances_偏向高基数特征和有更多分裂机会的特征。不要直接用它下结论，应配合第8章的SHAP分析进行交叉验证。', indent=True)
add_para('样本量不足的处理：平谷区仅约65个1km网格，对于500棵树的XGBoost可能不够。解决方案包括：从周边区县扩展更多网格（顺义、密云、三河）；使用CMIP6未来气候数据进行数据增强（将同一位置的不同年份/情景视为独立样本）。', indent=True)

add_heading_styled('7.6  掌握后的延展能力', 3)
add_para('XGBoost是结构化数据（表格数据）建模的王者，2015-2023年间赢得了绝大多数Kaggle竞赛。掌握它后，你可以：直接参与任何表格数据相关的机器学习竞赛或项目；理解几乎所有集成学习方法的原理（LightGBM、CatBoost、随机森林）；为学习深度学习打下基础（梯度提升中的"梯度"与神经网络反向传播中的"梯度"在数学上是同一个概念）。当前趋势（2025-2026）：XGBoost仍然是表格数据的首选模型，但神经网络的TabTransformer和FT-Transformer等架构正在追赶。你在这里学到的梯度提升原理是永恒的，无论工具如何演进。', indent=True)

doc.add_page_break()

# ── Chapter 8 ──
add_heading_styled('第8章  模型解释：SHAP值原理与应用', 2)
add_para('核心库：shap, matplotlib, seaborn', indent=True)

add_heading_styled('8.1  你能学到什么', 3)
add_para('本章教你"打开黑箱"：理解机器学习模型为什么做出某个预测。你将学会：Shapley值的博弈论起源及其公平分配原理；使用SHAP的TreeExplainer高效解释XGBoost模型；绘制Summary Plot（蜂群图）和Dependence Plot（依赖图）；计算特征交互效应。', indent=True)

add_heading_styled('8.2  Shapley值的直觉理解', 3)
add_para('Shapley值来自合作博弈论（Lloyd Shapley, 1953年，2012年诺贝尔经济学奖）。核心问题：一个团队完成了一个项目获得了一笔奖金，如何公平地将奖金分配给每个成员？公平的定义是：每个人的份额应等于他们加入团队时带来的边际贡献。', indent=True)
add_para('SHAP（SHapley Additive exPlanations）将这个想法移植到机器学习中："团队"是所有输入特征，"奖金"是模型的预测值，"每个人"是单个特征。SHAP值回答：特征X对这个预测值的贡献是多少？', indent=True)
add_para('例如：模型预测某网格的小麦产量为6.2吨/公顷（而所有网格的平均产量是5.5吨/公顷）。SHAP分解：年均温10.5度C贡献了+0.3吨（正面），土壤有机碳1.2%贡献了+0.5吨（正面），年降水620mm贡献了-0.1吨（略负面，可能过湿）。总和：5.5 + 0.3 + 0.5 - 0.1 = 6.2。这就是可解释的AI。', indent=True)

add_heading_styled('8.3  核心代码', 3)
add_code('import shap')
add_code('')
add_code('# 使用TreeExplainer获取SHAP值')
add_code('explainer = shap.TreeExplainer(model)')
add_code('shap_values = explainer.shap_values(X_test)')
add_code('')
add_code('# Summary Plot：全局特征重要性')
add_code('shap.summary_plot(shap_values, X_test, feature_names=feature_names)')
add_code('')
add_code('# Dependence Plot：单个特征与SHAP值的关系')
add_code('shap.dependence_plot("bio1", shap_values, X_test)')
add_code('')
add_code('# 单个预测的瀑布图（解释一个特定网格的预测）')
add_code('shap.waterfall_plot(explainer.expected_value,')
add_code('                    shap_values[0], X_test.iloc[0])')

add_heading_styled('8.4  SHAP在育种中的应用', 3)
add_para('本项目中SHAP的独特应用是"性状缺口分析"（Trait Gap Analysis）：在品种性状特征中，哪些性状对产量缺口的贡献最大？如果SHAP分析显示"千粒重"对产量的边际贡献为+0.5吨/公顷（在现有遗传背景下），而审定的新品种最高千粒重可提高5克，那么可以预期产量增益约为0.5吨/公顷——这就是育种方向的定量预测。', indent=True)

add_heading_styled('8.5  掌握后的延展能力', 3)
add_para('可解释AI（XAI）是2025-2026年AI领域最热门的子方向之一。欧盟AI法案（EU AI Act, 2024年通过）要求高风险AI系统必须提供可解释性。掌握SHAP后，你可以在任何领域产出可信赖的模型解释：金融信贷决策（为什么拒绝这笔贷款）、医疗诊断（为什么判断为高风险）、政策评估（哪些因素驱动了政策效果）。当前趋势是SHAP与LLM（大语言模型）的结合——让LLM用自然语言描述SHAP分析结果，形成"自动生成的模型分析报告"。', indent=True)

doc.add_page_break()

# ── Chapter 9 ──
add_heading_styled('第9章  多目标优化：NSGA-II', 2)
add_para('核心库：pymoo', indent=True)

add_heading_styled('9.1  你能学到什么', 3)
add_para('本章介绍如何使用遗传算法解决多目标优化问题。你将学会：理解帕累托最优（Pareto Optimality）的概念；掌握NSGA-II算法的工作原理（非支配排序+拥挤度距离）；使用pymoo库定义优化问题并求解；理解约束优化与无约束优化的区别。', indent=True)

add_heading_styled('9.2  为什么多目标优化在农业中如此重要', 3)
add_para('农业生产本质上是一个多目标权衡问题。农民希望同时实现：产量最大化（经济目标），氮肥用量最小化（成本目标），氮淋溶和水消耗最小化（环境目标）。这些目标往往是冲突的——多施氮肥能增产，但增加了淋溶风险和环境成本。单目标优化（如"最大化产量"）给出的答案是施肥无限大——这毫无实际意义。多目标优化找到的是帕累托前沿：在所有可能的方案中，无法在不损害至少一个目标的情况下改善另一个目标的方案的集合。', indent=True)

add_heading_styled('9.3  NSGA-II算法原理', 3)
add_para('NSGA-II（Non-dominated Sorting Genetic Algorithm II, Deb et al. 2002）是目前最广泛使用的多目标进化算法。它的工作流程模拟了自然选择：', indent=True)
add_para('初始化：随机生成N个候选管理方案（种群），每个方案是一个参数向量（如[氮肥量, 灌溉量, 秸秆还田率]）。', indent=True)
add_para('评估：用训练好的XGBoost模型预测每个方案的产量和环境影响（代理评估，比跑APSIM快10000倍）。', indent=True)
add_para('选择：非支配排序将所有方案分为多个前沿层级。第一前沿（Pareto Front）中的方案彼此非支配——没有一个方案在所有目标上都优于另一个。在同一个前沿内，拥挤度距离（Crowding Distance）鼓励选择与邻居差异大的方案，以保持解的多样性。', indent=True)
add_para('变异与交叉：选中的优秀方案通过交叉（交换参数）和变异（随机微调参数）产生下一代方案。', indent=True)
add_para('迭代：重复"评估-选择-交叉-变异"数百代，最终收敛到近似的帕累托前沿。', indent=True)

add_heading_styled('9.4  核心代码', 3)
add_code('import numpy as np')
add_code('from pymoo.algorithms.moo.nsga2 import NSGA2')
add_code('from pymoo.core.problem import Problem')
add_code('from pymoo.optimize import minimize')
add_code('from pymoo.operators.crossover.sbx import SBX')
add_code('from pymoo.operators.mutation.pm import PM')
add_code('from pymoo.operators.sampling.rnd import FloatRandomSampling')
add_code('')
add_code('class FarmManagementProblem(Problem):')
add_code('    """农田管理多目标优化问题"""')
add_code('    def __init__(self, xgb_model, env_features):')
add_code('        # 3个决策变量：氮(kg/ha)、灌溉(mm)、秸秆还田(%)')
add_code('        super().__init__(')
add_code('            n_var=3, n_obj=2, n_constr=0,')
add_code('            xl=np.array([50, 0, 0]),    # 下界')
add_code('            xu=np.array([300, 600, 100])  # 上界')
add_code('        )')
add_code('        self.model = xgb_model')
add_code('        self.env = env_features')
add_code('')
add_code('    def _evaluate(self, x, out, *args, **kwargs):')
add_code('        # x: (n_population, 3) 管理参数矩阵')
add_code('        # 将管理参数与环境特征拼接为完整输入')
add_code('        n = x.shape[0]')
add_code('        X_full = np.hstack([')
add_code('            np.tile(self.env, (n, 1)),  # 环境特征')
add_code('            x                         # 管理参数')
add_code('        ])')
add_code('        preds = self.model.predict(X_full)')
add_code('        # 目标1: 产量 (最大化 -> 取负值用于最小化)')
add_code('        # 目标2: 氮淋溶 (最小化)')
add_code('        out["F"] = np.column_stack([-preds[:, 0], preds[:, 1]])')
add_code('')
add_code('problem = FarmManagementProblem(xgb_model, env_sample)')
add_code('algorithm = NSGA2(')
add_code('    pop_size=100,')
add_code('    sampling=FloatRandomSampling(),')
add_code('    crossover=SBX(prob=0.9, eta=15),')
add_code('    mutation=PM(prob=1/3, eta=20),')
add_code(')')
add_code('res = minimize(problem, algorithm, ("n_gen", 200))')
add_code('print(f"帕累托前沿方案数: {len(res.F)}")')

add_heading_styled('9.5  NSGA-II的局限性及本项目中的应对', 3)
add_para('直接使用NSGA-II搜索66维输入空间（~60维环境指纹 + 3维管理参数 + ~3维品种性状）是不可行的——"维数灾难"会使进化算法在可行时间内无法收敛。Xiao et al. (2024)的解决方案是先训练XGBoost模型作为APSIM的代理（surrogate），然后只对管理参数进行优化（环境指纹固定为特定位置的观测值）。我们采用相同策略：对于每个目标网格点，固定其环境指纹和品种性状，只搜索最优管理参数。', indent=True)
add_para('备选方案：如果NSGA-II的代理模型方法仍然过于复杂，可以退化为使用论文已发表的优化管理方案作为标签，训练一个MultiOutputRegressor直接预测最优管理参数。这是一种更简单的监督学习方法，适合初期快速出结果。', indent=True)

add_heading_styled('9.6  掌握后的延展能力', 3)
add_para('多目标优化是工程设计和决策支持的核心技术。掌握NSGA-II后，你可以解决任何需要"在多个冲突目标中找到最佳折衷"的问题：供应链优化（成本 vs 时效 vs 碳排放）、建筑能效设计（建设成本 vs 运营能耗）、药物分子设计（药效 vs 毒性 vs 合成难度）。当前趋势（2025-2026）是基于强化学习的多目标优化——使用神经网络替代进化算法，能更高效地探索高维决策空间。你在这里学到的帕累托最优概念，同样是强化学习的多目标版本（Multi-Objective RL）的基础。', indent=True)

doc.add_page_break()

# ── Chapter 10 ──
add_heading_styled('第10章  不确定性量化：Bootstrap方法', 2)
add_para('核心库：numpy, scipy, scikit-learn', indent=True)

add_heading_styled('10.1  你能学到什么', 3)
add_para('本章教你量化和表达模型预测的不确定性。你将学会：理解偶然不确定性(aleatoric)与认知不确定性(epistemic)的区别；使用Bootstrap重采样估计预测的置信区间；理解为什么"点估计"（只给出一个数字）在科研和决策中是不够的。', indent=True)

add_heading_styled('10.2  为什么不确定性量化至关重要', 3)
add_para('假设模型预测某网格的最优施氮量为180 kg/ha，产量为7.2吨/ha。如果是你自己家的地，你会严格按照180 kg/ha施肥吗？你一定会问：这个预测有多可靠？±5 kg/ha还是±50 kg/ha？如果有50%的概率实际产量在6.0-6.5吨之间，你可能会选择更保守的方案。这就是不确定性量化的意义——它为决策提供风险信息。', indent=True)
add_para('在论文中，报告预测值而不报告置信区间是不完整的。审稿人会问："你的模型预测的不确定性有多大？不同CMIP6气候情景下的预测范围是多少？"提前做好不确定性量化，不仅让你得出更稳健的结论，也让论文更有说服力。', indent=True)

add_heading_styled('10.3  Bootstrap原理', 3)
add_para('Bootstrap（自助法，Efron 1979）是统计学中最优雅的想法之一。核心思路：我们只有一份样本（n=65个网格），但我们想知道如果重新采样，模型的预测会如何变化。Bootstrap的解决方案是："假装"手中的样本就是总体，从中有放回地重采样n次，形成一个新的Bootstrap样本（有些网格被抽到多次，有些一次也没有），在这个新样本上重新训练模型并做出预测。重复这个过程1000次，就得到了1000个可能的世界中的预测值——这些值的分布就是不确定性。', indent=True)

add_heading_styled('10.4  核心代码', 3)
add_code('import numpy as np')
add_code('from sklearn.utils import resample')
add_code('')
add_code('def bootstrap_predictions(model, X, y, X_pred, n_bootstrap=1000):')
add_code('    """Bootstrap预测：生成预测值的分布"""')
add_code('    preds = np.zeros((n_bootstrap, X_pred.shape[0]))')
add_code('')
add_code('    for i in range(n_bootstrap):')
add_code('        # 有放回重采样训练数据')
add_code('        X_bs, y_bs = resample(X, y, n_samples=len(X))')
add_code('')
add_code('        # 在Bootstrap样本上训练')
add_code('        model.fit(X_bs, y_bs)')
add_code('')
add_code('        # 对目标位置做预测')
add_code('        preds[i] = model.predict(X_pred)')
add_code('')
add_code('    # 计算统计量')
add_code('    mean_pred = preds.mean(axis=0)')
add_code('    lower_ci = np.percentile(preds, 2.5, axis=0)  # 95% CI下界')
add_code('    upper_ci = np.percentile(preds, 97.5, axis=0) # 95% CI上界')
add_code('')
add_code('    return mean_pred, lower_ci, upper_ci')

add_heading_styled('10.5  不确定性分解', 3)
add_para('总预测不确定性可以分解为两个来源：', indent=True)
add_para('数据不确定性（偶然不确定性）：由测量误差、空间变异等引起。例如同一网格内不同年份的产量波动。这种不确定性是固有的，无法通过增加数据来消除。', indent=True)
add_para('模型不确定性（认知不确定性）：由训练数据不足、模型结构选择等引起。例如只有65个训练样本时，模型对不同超参数组合给出不同的预测。这种不确定性可以通过增加数据、改进模型来降低。', indent=True)
add_para('在论文中，建议将这两种不确定性分别报告，以体现科学严谨性。Bootstrap方法本身主要量化模型不确定性。数据不确定性可以额外通过XGBoost的quantile regression或Monte Carlo Dropout的变体来估计。', indent=True)

add_heading_styled('10.6  掌握后的延展能力', 3)
add_para('不确定性量化是所有科学计算和工程决策的基础。掌握了Bootstrap和置信区间后，你可以：在任何统计建模中提供误差棒和置信区间；理解贝叶斯统计的基本思想（Bootstrap是频率学派的方法，贝叶斯用后验分布表达不确定性，两者在概念上互补）；参与需要风险评估的场景（临床试验、金融风控、结构工程设计）。当前趋势（2025-2026）是"符合性预测"（Conformal Prediction）的兴起——它提供分布无关(distribution-free)的预测区间保证，比Bootstrap更适用于小样本深度学习模型。', indent=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# PART 4: R LANGUAGE
# ═══════════════════════════════════════════════════════════════
add_heading_styled('第四部分  R语言统计与可视化', 1)

# ── Chapter 11 ──
add_heading_styled('第11章  R语言统计分析', 2)
add_para('核心包：tidyverse (dplyr, tidyr, ggplot2), caret, randomForest, vegan', indent=True)

add_heading_styled('11.1  你能学到什么', 3)
add_para('本章教你使用R语言进行统计分析和建模。你将学会：R的tidyverse数据操作范式（管道操作符%>%）；使用R进行方差分析(ANOVA)、多元回归、混合效应模型等经典统计；理解何时用R优于Python（探索性数据分析、统计检验、学术图表）。', indent=True)

add_heading_styled('11.2  为什么项目需要R', 3)
add_para('Python在机器学习和工程方面更强，但R在统计推断和学术出版方面有独特优势：R的统计检验报告比Python更完备（效应量、假设检验诊断图自动生成）；ggplot2公认是科技论文出版质量的唯一标准绘图系统；R的生态学/农学专用包（vegan用于群落分析，agricolae用于田间试验设计，metan用于多环境试验分析）在Python中没有等价物。', indent=True)
add_para('在本项目中，R的角色是Python的互补，而非替代。Python处理大规模数据提取、模型训练、优化搜索；R处理统计分析、假设检验、论文图表制作。两者通过CSV文件在文件系统层面交互。', indent=True)

add_heading_styled('11.3  R与Python的互操作', 3)
add_para('现代数据科学工作流中，R和Python并不冲突。以下三种方式可以实现互操作：', indent=True)
add_para('方式一（最简单）：通过CSV文件传递。Python写出结果到CSV，R读入做统计和绘图。这是本项目推荐的方式。', indent=True)
add_para('方式二：通过reticulate包在R中调用Python。在R脚本中import Python模块、调用Python函数。适合需要频繁交互的场景。', indent=True)
add_para('方式三：通过Quarto或R Markdown。在同一个.qmd文件中混合编写Python和R代码块，生成统一的HTML/PDF报告。Quarto是RStudio开发的下一代科学出版系统，推荐用于最终报告生成。', indent=True)

add_heading_styled('11.4  核心操作示例', 3)
add_para('使用tidyverse进行数据操作（管道流）：', indent=True)
add_code('library(tidyverse)')
add_code('')
add_code('# 读取Python输出的特征矩阵')
add_code('df <- read_csv("features_matrix.csv")')
add_code('')
add_code('# 管道操作：筛选 -> 选择列 -> 分组 -> 汇总')
add_code('df %>%')
add_code('  filter(year == 2020) %>%')
add_code('  select(grid_id, bio1, bio12, yield) %>%')
add_code('  group_by(cut(bio12, breaks = 5)) %>%  # 按降水五分位分组')
add_code('  summarise(')
add_code('    mean_yield = mean(yield, na.rm = TRUE),')
add_code('    sd_yield = sd(yield, na.rm = TRUE),')
add_code('    n = n()')
add_code('  )')
add_para('方差分析（ANOVA）检验不同管理方案之间的产量差异是否显著：', indent=True)
add_code('# 单因素ANOVA')
add_code('model <- aov(yield ~ nitrogen_level, data = df)')
add_code('summary(model)')
add_code('TukeyHSD(model)  # 事后两两比较')
add_code('')
add_code('# 双因素ANOVA（氮水平 x 灌溉水平交互效应）')
add_code('model2 <- aov(yield ~ nitrogen_level * irrigation_level, data = df)')
add_code('summary(model2)')

add_heading_styled('11.5  掌握后的延展能力', 3)
add_para('R语言是统计学的通用语言，在以下领域有不可替代的作用：生物统计学（临床试验、流行病学）、生态学（群落分析、物种分布模型）、数量遗传学（GWAS、基因组选择）、社会科学（结构方程模型、社会网络分析）。当前趋势（2025-2026）是Posit（原RStudio）公司推动的"多语言数据科学平台"——在同一个IDE中同时使用R、Python、Julia和SQL。你在这里掌握的双语言能力正是这一趋势的核心竞争力。', indent=True)

doc.add_page_break()

# ── Chapter 12 ──
add_heading_styled('第12章  R语言科学可视化', 2)
add_para('核心包：ggplot2, patchwork, ggsci, ggrepel, sf, tmap', indent=True)

add_heading_styled('12.1  你能学到什么', 3)
add_para('本章教你使用ggplot2制作出版级别的科学图表。你将学会：ggplot2的分层语法（Grammar of Graphics）——将图表分解为数据、几何对象、坐标系、分面和主题五个独立层次；制作论文中常见的图表类型：箱线图、散点图矩阵、热图、地图；使用patchwork组合多张图为一张复合图；使用ggsci应用学术期刊的配色方案（Nature、Science、Lancet风格）。', indent=True)

add_heading_styled('12.2  ggplot2语法哲学', 3)
add_para('ggplot2的核心思想是"图形语法"（Grammar of Graphics, Wilkinson 2005）：任何统计图形都可以分解为以下组件的组合：data（要可视化的数据框）、aesthetics（美学映射——哪个变量映射到x轴、y轴、颜色、大小）、geometries（几何对象——点、线、柱、箱）、facets（分面——按分类变量拆分为子图）、themes（主题——字体、背景、网格线）。这种分层设计使得ggplot2比matplotlib更易于构建复杂图表：你不需要计算每个点的像素坐标，只需要声明变量之间的映射关系。', indent=True)

add_heading_styled('12.3  核心代码示例', 3)
add_code('library(ggplot2)')
add_code('library(patchwork)')
add_code('library(ggsci)')
add_code('')
add_code('# 图1: SHAP特征重要性（蜂群图用geom_point替代）')
add_code('p1 <- ggplot(shap_df, aes(x = reorder(feature, abs_shap), y = shap_value, color = feature_value)) +')
add_code('  geom_jitter(alpha = 0.6, width = 0.2) +')
add_code('  coord_flip() +')
add_code('  scale_color_gradient2(low = "blue", mid = "grey", high = "red") +')
add_code('  labs(x = "", y = "SHAP value", title = "特征重要性（SHAP）") +')
add_code('  theme_minimal()')
add_code('')
add_code('# 图2: 帕累托前沿散点图')
add_code('p2 <- ggplot(pareto_df, aes(x = yield, y = n_leaching, color = irrigation)) +')
add_code('  geom_point(size = 3, alpha = 0.8) +')
add_code('  scale_color_viridis_c() +')
add_code('  labs(x = "产量 (t/ha)", y = "氮淋溶 (kg/ha)",')
add_code('       title = "帕累托前沿：产量 vs 环境代价") +')
add_code('  theme_classic()')
add_code('')
add_code('# 图3: 不确定性带（Bootstrap CI）')
add_code('p3 <- ggplot(ci_df, aes(x = n_level, y = yield_pred)) +')
add_code('  geom_ribbon(aes(ymin = ci_lower, ymax = ci_upper), alpha = 0.3) +')
add_code('  geom_line(size = 1, color = "steelblue") +')
add_code('  labs(x = "施氮量 (kg/ha)", y = "预测产量 (t/ha)")')
add_code('')
add_code('# 使用patchwork组合三张图')
add_code('combined <- (p1 | p2) / p3 +')
add_code('  plot_annotation(tag_levels = "a")  # 自动添加(a)(b)(c)标签')
add_code('ggsave("figure_combined.pdf", combined, width = 12, height = 8)')

add_heading_styled('12.4  论文图表规范', 3)
add_para('制作论文图表时需遵循以下规范：所有文字使用Times New Roman字体（与正文一致）；配色使用色盲友好的调色板（viridis / cividis / ggsci::scale_color_npg）；字号不小于8pt，确保印刷后可读；图表标签使用(a)(b)(c)而非1,2,3（遵循大多数期刊规范）；导出为PDF或600 DPI的TIFF格式（而非PNG/JPG）。', indent=True)

add_heading_styled('12.5  掌握后的延展能力', 3)
add_para('ggplot2是数据可视化领域的标准工具，掌握后你可以制作任何领域的出版级图表。当前趋势（2025-2026）有两个方向：交互式可视化（plotly、gganimate）将静态图表变为可交互的Web应用；AI辅助可视化——通过自然语言描述生成ggplot2代码（如ChatGPT的"生成一张按年份分组的箱线图"）。你在这里学到的图形语法思想，同样适用于Python的plotnine（ggplot2的Python移植）和Vega-Lite（声明式可视化的JSON规范）。', indent=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# PART 5: DEPLOYMENT
# ═══════════════════════════════════════════════════════════════
add_heading_styled('第五部分  部署与可重复性', 1)

# ── Chapter 13 ──
add_heading_styled('第13章  模型部署与API服务', 2)
add_para('核心库：FastAPI, uvicorn, pydantic', indent=True)

add_heading_styled('13.1  你能学到什么', 3)
add_para('本章教你将训练好的模型部署为可被他人调用的Web服务。你将学会：使用FastAPI构建RESTful API（当前Python最快的Web框架之一）；定义输入输出的数据模式（Pydantic模型）；使用uvicorn启动生产级服务器；编写API文档（FastAPI自动生成Swagger UI）。', indent=True)

add_heading_styled('13.2  为什么需要部署', 3)
add_para('训练好的模型如果只存在于你的笔记本中，它只是一堆数字。只有部署为API（应用程序编程接口），其他人才能在不安装Python环境的情况下使用你的模型：通过浏览器或一个简单的HTTP请求，输入环境条件，获得管理建议和产量预测。部署也让你的工作可以被引用和复现——审稿人可以通过API验证你的结果。', indent=True)

add_heading_styled('13.3  核心代码', 3)
add_code('from fastapi import FastAPI')
add_code('from pydantic import BaseModel')
add_code('import joblib')
add_code('import numpy as np')
add_code('')
add_code('app = FastAPI(title="2026-SP Management Optimizer")')
add_code('')
add_code('# 定义输入数据模式')
add_code('class EnvironmentalInput(BaseModel):')
add_code('    bio1: float    # 年均温')
add_code('    bio12: float   # 年降水')
add_code('    sand: float    # 砂粒含量')
add_code('    soc: float     # 土壤有机碳')
add_code('    elev: float    # 海拔')
add_code('    # ... 其余60个特征字段')
add_code('')
add_code('# 定义输出数据模式')
add_code('class ManagementOutput(BaseModel):')
add_code('    n_optimal: float       # 最优施氮量 (kg/ha)')
add_code('    irr_optimal: float     # 最优灌溉量 (mm)')
add_code('    yield_predicted: float # 预测产量 (t/ha)')
add_code('    yield_ci_lower: float  # 产量95%CI下界')
add_code('    yield_ci_upper: float  # 产量95%CI上界')
add_code('')
add_code('# 启动时加载模型')
add_code('@app.on_event("startup")')
add_code('async def load_models():')
add_code('    global model, scaler')
add_code('    model = joblib.load("model_xgb.pkl")')
add_code('    scaler = joblib.load("scaler.pkl")')
add_code('')
add_code('@app.post("/predict", response_model=ManagementOutput)')
add_code('async def predict(input_data: EnvironmentalInput):')
add_code('    features = np.array([[input_data.bio1, input_data.bio12, ...]])')
add_code('    features_scaled = scaler.transform(features)')
add_code('    pred = model.predict(features_scaled)[0]')
add_code('    return ManagementOutput(')
add_code('        n_optimal=pred[0], irr_optimal=pred[1],')
add_code('        yield_predicted=pred[2],')
add_code('        yield_ci_lower=pred[3], yield_ci_upper=pred[4]')
add_code('    )')
add_para('启动服务：在命令行运行uvicorn app:app --host 0.0.0.0 --port 8000，然后打开浏览器访问 http://localhost:8000/docs ，即可看到自动生成的交互式API文档（Swagger UI），可以直接在网页上测试API。', indent=True)

add_heading_styled('13.4  掌握后的延展能力', 3)
add_para('API部署是所有AI产品从"实验室模型"走向"实际应用"的关键一步。掌握FastAPI后，你可以：将任何Python函数包装为可供Web/App调用的微服务；参与云原生架构的设计与开发；理解后端开发的基本范式（请求-响应模型、中间件、异步处理）。当前趋势（2025-2026）是"MLOps"——机器学习运维，将模型训练-部署-监控形成闭环。你在这里学到的API部署是MLOps链路中"部署"环节的核心技能。', indent=True)

doc.add_page_break()

# ── Chapter 14 ──
add_heading_styled('第14章  容器化与可重复研究', 2)
add_para('核心工具：Docker', indent=True)

add_heading_styled('14.1  你能学到什么', 3)
add_para('本章教你使用Docker打包整个项目的运行环境。你将学会：编写Dockerfile定义项目的完整环境；构建Docker镜像并推送到Docker Hub；使用docker-compose管理多服务编排；理解容器化如何确保研究的可重复性。', indent=True)

add_heading_styled('14.2  "在我的机器上能跑"的问题', 3)
add_para('这是科学计算最常见的困境：你的代码在你的Windows笔记本电脑上完美运行，但coworker的Mac上报错"找不到geos_c.dll"；审稿人的Linux服务器上报错"GLIBC版本不匹配"；你自己的电脑重装系统后环境全部丢失。Docker通过将应用程序及其所有依赖（包括操作系统级库）打包为一个轻量级的"容器"，确保在任何机器上运行结果完全一致。', indent=True)

add_heading_styled('14.3  项目Dockerfile', 3)
add_code('FROM python:3.11-slim')
add_code('')
add_code('# 安装系统级依赖（GDAL、proj等地理空间库）')
add_code('RUN apt-get update && apt-get install -y \\')
add_code('    gdal-bin libgdal-dev libproj-dev && \\')
add_code('    rm -rf /var/lib/apt/lists/*')
add_code('')
add_code('# 设置工作目录')
add_code('WORKDIR /app')
add_code('')
add_code('# 复制依赖文件并安装')
add_code('COPY requirements.txt .')
add_code('RUN pip install --no-cache-dir -r requirements.txt')
add_code('')
add_code('# 复制源代码')
add_code('COPY src/ ./src/')
add_code('COPY config.yaml .')
add_code('')
add_code('# 暴露API端口')
add_code('EXPOSE 8000')
add_code('')
add_code('# 启动命令')
add_code('CMD ["uvicorn", "src.api.main:app", "--host", "0.0.0.0", "--port", "8000"]')

add_heading_styled('14.4  可重复研究的最佳实践', 3)
add_para('除了Docker容器化，可重复研究还需要以下配套措施：', indent=True)
add_para('随机种子固定：在所有涉及随机性的操作中设置固定种子（numpy.random.seed(42), xgboost的random_state参数, sklearn的random_state）。这是可重复性的最基本要求。', indent=True)
add_para('数据版本化：使用DVC（Data Version Control）或直接下载日期标注原始数据，确保所有合作者使用完全相同的数据版本。', indent=True)
add_para('配置外部化：所有可调参数（文件路径、模型超参数、网格分辨率）抽取到config.yaml文件中，禁止硬编码在Python代码中。修改参数只需编辑配置文件，无需改动代码。', indent=True)
add_para('日志记录：使用Python的logging模块（而非print）记录训练过程，包括时间戳、损失值、验证分数。这对于后续排查"为什么上次训练效果好而这次不好"至关重要。', indent=True)

add_heading_styled('14.5  掌握后的延展能力', 3)
add_para('容器化是现代软件工程和云原生架构的基础。掌握Docker后，你可以：将任何应用部署到Kubernetes集群进行自动扩缩容；使用GitHub Actions自动构建和推送镜像（CI/CD）；参与微服务架构的设计。当前趋势（2025-2026）是"WebAssembly"（Wasm）作为轻量级容器替代方案的兴起——它比Docker启动更快（毫秒级），安全性更高，但在生态成熟度上仍落后于Docker。', indent=True)

doc.add_page_break()

# ── Chapter 15 ──
add_heading_styled('第15章  学术写作与文献管理', 2)

add_heading_styled('15.1  你能学到什么', 3)
add_para('本章介绍学术论文写作和文献管理的工具与实践。你将学会：使用Zotero管理参考文献并自动生成格式化引用；遵循学术写作的结构规范（IMRaD格式）；使用Markdown或LaTeX撰写论文初稿；在写作过程中保持"去AI化"——使文字具有个人风格而非明显的AI生成痕迹。', indent=True)

add_heading_styled('15.2  文献管理：Zotero', 3)
add_para('Zotero是免费开源的文献管理工具，推荐原因：浏览器插件一键抓取论文元数据（标题、作者、期刊、DOI）；自动重命名和整理PDF文件；Word/LibreOffice插件支持"引用-刷新"即时生成参考文献列表；支持共享文献库（团队协作时所有人使用同一套参考文献）。', indent=True)
add_para('本项目的核心参考文献及其Zotero标签建议：Xiao et al. 2024 (Nature Food) — 标签"method-benchmark"（方法对标）；Deb et al. 2002 (IEEE Trans Evol Comput) — 标签"method-nsga2"（NSGA-II原论文）；Chen & Guestrin 2016 (KDD) — 标签"method-xgboost"（XGBoost原论文）；Lundberg & Lee 2017 (NIPS) — 标签"method-shap"（SHAP原论文）；WorldClim相关文献 — 标签"data-worldclim"；SoilGrids相关文献 — 标签"data-soil"。', indent=True)

add_heading_styled('15.3  论文结构设计', 3)
add_para('推荐使用IMRaD结构（Introduction, Methods, Results, and Discussion），这是绝大多数SCI期刊的标准格式。以下为本项目的具体章节设计：', indent=True)
add_table(
    ['章节', '内容', '预计字数'],
    [
        ['Introduction', '研究背景、华北平原小麦-玉米生产的环境挑战、前人工作(Xiao et al. 2024)、本研究的目标与创新点', '~800词'],
        ['Materials and Methods', '2.1 研究区域(平谷区); 2.2 数据来源(逐源列表); 2.3 环境指纹构建方法; 2.4 XGBoost模型训练与验证; 2.5 NSGA-II优化; 2.6 SHAP解释; 2.7 Bootstrap不确定性', '~1500词'],
        ['Results', '3.1 模型性能; 3.2 优化管理方案 vs 农户实践; 3.3 SHAP特征重要性分析; 3.4 育种方向预测; 3.5 不确定性分析', '~1200词'],
        ['Discussion', '4.1 管理优化方案的可操作性; 4.2 1km分辨率的有效性; 4.3 品种数据的局限性; 4.4 与Xiao et al.的对比; 4.5 未来方向', '~1000词'],
        ['Conclusion', '简明总结主要发现和意义', '~200词'],
    ],
    [3, 10, 3.5]
)

add_heading_styled('15.4  写作规范与去AI化', 3)
add_para('学术论文写作中需特别注意以下规范：全文使用被动语态为主（"The model was trained"而非"We trained the model"），部分期刊允许第一人称但需保持一致；数据和单位之间留空格（"180 kg/ha"而非"180kg/ha"）；所有缩写首次出现时给出全称（"Growing Degree Days (GDD)"）；参考文献格式严格遵循目标期刊的要求（Journal of Integrative Agriculture使用APA格式）。', indent=True)
add_para('关于去AI化：AI辅助写作是合理的，但直接使用AI生成且不经修改的文字容易被识别。以下方法可以降低AI痕迹：将长句拆分为短句（AI喜欢嵌套从句）；使用你自己熟悉的表达方式替换AI的套话（如将"In the context of"替换为具体的情境描述）；主动使用领域内的术语而非AI喜欢的通用表达（如"粒重"而非"籽粒重量参数"）；在Discussion中引用你自己阅读过的文献的具体结论，而非AI生成的泛泛讨论。', indent=True)

add_heading_styled('15.5  掌握后的延展能力', 3)
add_para('学术写作是所有研究者的基本功。掌握了IMRaD结构和文献管理后，你可以高效地撰写任何学科的学术论文、学位论文和基金申请书。当前趋势（2025-2026）是"动态文档"（Dynamic Documents）——使用Quarto或Jupyter Book将代码、图表和文本整合为可自动更新的出版级文档。掌握这一技能意味着你的论文图表可以随着数据更新而自动刷新，无需手动重新截图替换。', indent=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# APPENDICES
# ═══════════════════════════════════════════════════════════════
add_heading_styled('附录', 1)

# ── Appendix A ──
add_heading_styled('附录A  推荐学习路径（3周速成方案）', 2)

add_table(
    ['时间', '学习内容', '产出'],
    [
        ['第1周', 'Git基础 + Python环境搭建 + pandas核心操作', '能独立clone仓库、创建分支、用pandas处理表格数据'],
        ['第2周', 'rasterio/geopandas + XGBoost基础 + SHAP初步', '能提取栅格数据、训练简单的XGBoost模型并解释特征重要性'],
        ['第3周', 'NSGA-II/遗传算法 + R语言ggplot2 + FastAPI基础', '能运行多目标优化、制作出版级图表、部署简单的API'],
    ],
    [2.5, 9.5, 5.5]
)

# ── Appendix B ──
add_heading_styled('附录B  常见问题排查', 2)

faqs = [
    ('rasterio导入报错"DLL load failed"',
     '原因：Windows缺少GDAL的DLL依赖。解决：使用conda安装而非pip（conda install -c conda-forge rasterio），conda会自动处理系统级依赖。'),
    ('XGBoost训练报错"Check failed: valid"',
     '原因：数据中包含NaN或无穷值。解决：训练前检查df.isnull().sum()和np.isinf(X).sum()。'),
    ('Git push被拒绝"Updates were rejected"',
     '原因：远程分支有新的提交，你本地的分支落后了。解决：git pull --rebase origin main，解决冲突后再push。'),
    ('NSGA-II收敛慢或未收敛',
     '原因：决策变量范围过大或种群数量太小。解决：增大pop_size到200+，增加n_gen到500+，确保每个决策变量的上下界合理。'),
    ('ggplot2中文字体显示为方块',
     '原因：R未配置中文字体。解决：安装showtext包，使用showtext_auto()启用，然后用font_add()注册中文字体。'),
    ('SHAP summary_plot报错"matplotlib error"',
     '原因：SHAP默认绘图与某些matplotlib版本不兼容。解决：升级到shap>=0.42和matplotlib>=3.7，或使用shap.plots.beeswarm代替。'),
]

for q, a in faqs:
    add_para(f'Q: {q}', bold=True, indent=True)
    add_para(f'A: {a}', indent=True)

# ── Appendix C ──
doc.add_page_break()
add_heading_styled('附录C  术语对照表', 2)

add_table(
    ['英文术语', '中文', '释义'],
    [
        ['XGBoost', '极限梯度提升', '基于梯度提升的集成学习算法，结构化数据建模的首选'],
        ['SHAP', 'Shapley加法解释', '基于博弈论Shapley值的模型解释方法'],
        ['NSGA-II', '非支配排序遗传算法II', '基于帕累托最优的多目标进化优化算法'],
        ['GDD', '生长度日', '作物发育所需的热量累积指标，单位为度C-日'],
        ['Bootstrap', '自助法', '通过有放回重采样估计统计量分布的方法'],
        ['Raster', '栅格数据', '由规则网格组成的空间数据，如GeoTIFF卫星影像'],
        ['CRS', '坐标参考系', '定义地理坐标与平面坐标转换关系的系统'],
        ['Affine Transform', '仿射变换', '描述栅格像素行列号与地理坐标之间的线性变换'],
        ['Feature Engineering', '特征工程', '将原始数据转化为模型可用的特征向量的过程'],
        ['Overfitting', '过拟合', '模型在训练集上表现极好但在新数据上表现差的现象'],
        ['Cross Validation', '交叉验证', '将数据分为K份，轮流用K-1份训练1份验证的评估方法'],
        ['Hyperparameter', '超参数', '在训练前设定的参数（如树深度、学习率），不能从数据中学习'],
        ['Pareto Front', '帕累托前沿', '多目标优化中所有非支配解的集合'],
        ['Surrogate Model', '代理模型', '用于替代复杂昂贵模拟的快速近似模型'],
        ['CI/CD', '持续集成/持续部署', '自动测试和部署代码的DevOps实践'],
        ['API', '应用程序编程接口', '软件组件之间通信的标准化接口'],
        ['Docker', '容器化平台', '将应用及其依赖环境打包为可移植容器的工具'],
        ['MLOps', '机器学习运维', '将机器学习模型从开发到部署的完整生命周期管理'],
    ],
    [4, 4, 9]
)

# ── End ──
doc.add_paragraph()
add_para('---', indent=True)
add_para('本手册版本 1.0 | 2026年7月 | 如有更新将在GitHub仓库中同步发布。', indent=True)
add_para('项目仓库：[待创建]', indent=True)
add_para('问题反馈：请在GitHub Issues中提交。', indent=True)

# Save
output_path = r'D:\2026-SP\育种模型项目技术培训手册.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)

# Report stats
import zipfile
size_kb = os.path.getsize(output_path) / 1024
with zipfile.ZipFile(output_path) as z:
    xml_size = sum(info.file_size for info in z.infolist() if info.filename.endswith('.xml'))
    
print(f'Saved: {output_path}')
print(f'File size: {size_kb:.0f} KB')
print(f'Paragraphs: {len(doc.paragraphs)}')
print(f'Tables: {len(doc.tables)}')
