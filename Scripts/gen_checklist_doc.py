#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Generate 数据清单 from CSV as clean Word table"""

import csv, os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

csv_path = r'D:\2026-SP\Data\data_checklist.csv'
doc_path = r'D:\2026-SP\数据获取状态清单.docx'

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
font.size = Pt(10)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(2)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(12)
run = title.add_run('育种模型环境指纹：数据获取状态清单')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# Summary
summary = doc.add_paragraph()
summary.paragraph_format.space_after = Pt(6)
run = summary.add_run('彭宇程 | 2026年7月19日 | 总维度数：65')
run.font.size = Pt(10)

# Read CSV
rows = []
with open(csv_path, 'r', encoding='utf-8-sig') as f:
    reader = csv.DictReader(f)
    for row in reader:
        rows.append(row)

# Count status
got = sum(1 for r in rows if '已获取' in r['状态'] or '可计算' in r['状态'])
manual = sum(1 for r in rows if '需手动' in r['状态'])

stats = doc.add_paragraph()
stats.paragraph_format.space_after = Pt(12)
run = stats.add_run(f'已获取/可计算：{got}维   需手动：{manual}维   完成度：{got}/65（{got*100//65}%）')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Create table
table = doc.add_table(rows=len(rows)+1, cols=5)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Headers
headers = ['序号', '特征名称', '含义', '本地文件', '状态']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    # Light gray header
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'E8E8E8')
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

# Data rows
for ri, row in enumerate(rows):
    vals = [row['维度序号'], row['特征名称'], row['含义'], row['本地文件'], row['状态']]
    for ci, val in enumerate(vals):
        cell = table.rows[ri+1].cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        if ci == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # Color status column
        if ci == 4:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if '已获取' in val:
                run.font.color.rgb = RGBColor(0x1B, 0x7A, 0x2B)  # green
            elif '可计算' in val:
                run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)  # blue
            elif '模板' in val:
                run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)  # orange
            elif '需手动' in val:
                run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)  # red
    
    # Alternate row shading
    if ri % 2 == 0:
        for ci in range(5):
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F5F5F5')
            shading.set(qn('w:val'), 'clear')
            table.rows[ri+1].cells[ci]._tc.get_or_add_tcPr().append(shading)

# Set column widths
widths = [Cm(1.2), Cm(2.8), Cm(5.0), Cm(6.2), Cm(2.2)]
for row in table.rows:
    for ci, w in enumerate(widths):
        row.cells[ci].width = w

# Legend
doc.add_paragraph()
legend = doc.add_paragraph()
legend.paragraph_format.space_before = Pt(6)
for text, color in [
    ('已获取', RGBColor(0x1B, 0x7A, 0x2B)),
    ('可计算', RGBColor(0x00, 0x66, 0xCC)),
    ('模板已有待填入', RGBColor(0xCC, 0x88, 0x00)),
    ('需手动', RGBColor(0xCC, 0x33, 0x33)),
]:
    run = legend.add_run(f'■ {text}  ')
    run.font.color.rgb = color
    run.font.size = Pt(9)

doc.save(doc_path)
print(f'Saved: {doc_path}')
print(f'Stats: {got} acquired/computable, {manual} manual, total {len(rows)} dims')
