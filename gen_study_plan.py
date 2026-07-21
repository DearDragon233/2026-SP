"""
30天雅思+生物英语学习计划生成
Word文档，A4，中文排版
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ===== 页面设置 =====
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ===== 样式设置 =====
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
r = style.element.rPr
if r is None:
    r = OxmlElement('w:rPr')
    style.element.append(r)
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:eastAsia'), '宋体')
r.append(rFonts)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        r = run._element.rPr
        if r is None:
            r = OxmlElement('w:rPr')
            run._element.append(r)
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '黑体')
        r.append(rFonts)
    return h

def add_para(text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def set_cell_text(cell, text, bold=False, size=10):
    for r in cell.paragraphs[0].runs:
        r.text = ''
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold

def make_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            set_cell_text(table.rows[ri+1].cells[ci], val, size=9)
    return table

# ===== 封面 =====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('30天雅思词汇 + 生物专业英语\n强化学习计划')
run.font.size = Pt(22)
run.font.name = 'Times New Roman'
run.bold = True

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('IELTS Vocabulary & Biology English\nIntensive 30-Day Program')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.italic = True

doc.add_paragraph()
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('目标：雅思6.5-7.0 + 生物专业文献阅读能力\n时间：30天，每日约2-3小时\n2026年7月')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_page_break()

# ===== 总体框架 =====
add_heading_styled('一、学习计划总体框架', 1)

add_para('本计划按30天设计，每天包含四项内容：雅思词汇、生物专业词汇、语法点、以及每3天一个完整的生物实验英文流程。')

make_table(
    ['模块', '频率', '每日量', '30天总量'],
    [
        ['雅思核心词汇', '每日', '30词', '900词'],
        ['生物专业词汇', '每日', '15词', '450词'],
        ['语法点', '每日', '1-2个', '45个'],
        ['英文生物实验流程', '每3天', '1个完整protocol', '10个'],
        ['复习/测验', '每6天', '1次综合', '5次'],
    ]
)

doc.add_paragraph()

# ===== 30天日程总表 =====
add_heading_styled('二、30天日程总览', 1)

days = [
    ('Day 1', 'IELTS学术词汇1-30', 'Cell/Molecular Biology基础(1-15)', '一般现在时 vs 现在进行时', '—'),
    ('Day 2', 'IELTS学术词汇31-60', 'Cell/Molecular Biology基础(16-30)', '一般过去时 vs 过去进行时', '—'),
    ('Day 3', 'IELTS学术词汇61-90', 'Cell/Molecular Biology基础(31-45)', '现在完成时 vs 过去完成时', 'Protocol 1: DNA Extraction'),
    ('Day 4', 'IELTS学术词汇91-120', 'Genetics术语(1-15)', '被动语态(一般时)', '—'),
    ('Day 5', 'IELTS学术词汇121-150', 'Genetics术语(16-30)', '被动语态(完成时)', '—'),
    ('Day 6', 'IELTS学术词汇151-180', 'Genetics术语(31-45)', '情态动词(can/could/may/might)', '复习测验1'),
    ('Day 7', 'IELTS学术词汇181-210', 'Breeding & Agriculture(1-15)', '情态动词(must/should/will/would)', '—'),
    ('Day 8', 'IELTS学术词汇211-240', 'Breeding & Agriculture(16-30)', '条件句(零条件/第一条件)', '—'),
    ('Day 9', 'IELTS学术词汇241-270', 'Breeding & Agriculture(31-45)', '条件句(第二/第三条件)', 'Protocol 2: PCR Amplification'),
    ('Day 10', 'IELTS学术词汇271-300', 'Biochemistry术语(1-15)', '关系从句(who/which/that)', '—'),
    ('Day 11', 'IELTS环境类词汇(1-30)', 'Biochemistry术语(16-30)', '关系从句(where/when/whose)', '—'),
    ('Day 12', 'IELTS环境类词汇(31-60)', 'Biochemistry术语(31-45)', '不定式与动名词', '复习测验2'),
    ('Day 13', 'IELTS环境类词汇(61-90)', 'Plant Physiology(1-15)', '冠词(a/an/the)精讲', '—'),
    ('Day 14', 'IELTS科技类词汇(1-30)', 'Plant Physiology(16-30)', '可数/不可数名词', '—'),
    ('Day 15', 'IELTS科技类词汇(31-60)', 'Plant Physiology(31-45)', '比较级与最高级', 'Protocol 3: Gel Electrophoresis'),
    ('Day 16', 'IELTS科技类词汇(61-90)', 'Microbiology术语(1-15)', '间接引语(陈述句)', '—'),
    ('Day 17', 'IELTS健康类词汇(1-30)', 'Microbiology术语(16-30)', '间接引语(疑问句)', '—'),
    ('Day 18', 'IELTS健康类词汇(31-60)', 'Microbiology术语(31-45)', '虚拟语气(wish/if only)', '复习测验3'),
    ('Day 19', 'IELTS健康类词汇(61-90)', 'Statistics & Data(1-15)', '倒装句', '—'),
    ('Day 20', 'IELTS教育类词汇(1-30)', 'Statistics & Data(16-30)', '强调句(It is...that)', '—'),
    ('Day 21', 'IELTS教育类词汇(31-60)', 'Statistics & Data(31-45)', '并列连词与从属连词', 'Protocol 4: Bacterial Transformation'),
    ('Day 22', 'IELTS教育类词汇(61-90)', 'Bioinformatics(1-15)', '介词搭配(时间/地点)', '—'),
    ('Day 23', 'IELTS写作高频词汇(1-30)', 'Bioinformatics(16-30)', '介词搭配(动词+介词)', '—'),
    ('Day 24', 'IELTS写作高频词汇(31-60)', 'Bioinformatics(31-45)', '主谓一致', '复习测验4'),
    ('Day 25', 'IELTS写作高频词汇(61-90)', 'Ecology & Evolution(1-15)', '定语从句 vs 分词短语', '—'),
    ('Day 26', 'IELTS口语高频词汇(1-30)', 'Ecology & Evolution(16-30)', '名词性从句', '—'),
    ('Day 27', 'IELTS口语高频词汇(31-60)', 'Ecology & Evolution(31-45)', '状语从句全总结', 'Protocol 5: RNA Extraction'),
    ('Day 28', 'IELTS口语高频词汇(61-90)', 'Breeding Methods(1-15)', '学术写作常用句型(1)', '—'),
    ('Day 29', 'IELTS真题高频词汇(1-30)', 'Breeding Methods(16-30)', '学术写作常用句型(2)', '—'),
    ('Day 30', 'IELTS真题高频词汇(31-60)', 'Breeding Methods(31-45)', '语法总复习+易错点', '复习测验5 + 总复习'),
]

make_table(
    ['Day', '雅思词汇(30词/天)', '生物专业词汇(15词/天)', '语法点(1-2个/天)', '实验流程/测验'],
    [(d[0], d[1], d[2], d[3], d[4]) for d in days]
)

doc.add_page_break()

# ===== 雅思词汇 =====
add_heading_styled('三、雅思核心词汇表（900词）', 1)

ielts_vocab = {
    'Day 1-3: 学术高频词汇(1-90)': [
        'analyze, approach, area, assess, assume, authority, available, benefit, concept, consist',
        'constitute, context, contract, create, data, define, derive, distribute, economy, environment',
        'establish, estimate, evident, export, factor, finance, formula, function, identify, income',
        'indicate, individual, interpret, involve, issue, labour, legal, legislate, major, method',
        'occur, percent, period, policy, principle, proceed, process, require, research, respond',
        'role, section, sector, significant, similar, source, specific, structure, theory, vary',
        'achieve, acquire, adapt, adequate, adjacent, adjust, administrate, adult, advocate, affect',
        'aggregate, aid, albeit, allocate, alter, alternative, ambiguous, amend, analogy, annual',
        'anticipate, apparent, append, appreciate, approach, appropriate, approximate, arbitrary, aspect, assemble',
    ],
    'Day 4-6: 学术高频词汇(91-180)': [
        'assess, assign, assist, assume, assure, attach, attain, attitude, attribute, author',
        'automate, aware, behalf, benefit, bias, bond, brief, bulk, capable, capacity',
        'category, cease, challenge, channel, chapter, chart, chemical, circumstance, cite, civil',
        'clarify, classic, clause, code, coherent, coincide, collapse, colleague, commence, comment',
        'commission, commit, commodity, communicate, community, compatible, compensate, compile, complement, complex',
        'component, compound, comprehensive, comprise, compute, conceive, concentrate, concept, conclude, concurrent',
        'conduct, confer, confine, confirm, conflict, conform, consent, consequent, considerable, consist',
        'constant, constitute, constrain, construct, consult, consume, contact, contemporary, context, contradict',
        'contrary, contrast, contribute, controversy, convene, converse, convert, convince, cooperate, coordinate',
    ],
    'Day 7-10: 学术高频词汇(181-300)': [
        'core, corporate, correspond, couple, create, credit, criteria, crucial, culture, currency',
        'cycle, data, debate, decade, decline, deduce, define, definite, demonstrate, denote',
        'deny, depress, derive, design, despite, detect, deviate, device, devote, differentiate',
        'dimension, diminish, discrete, discriminate, displace, display, dispose, distinct, distort, distribute',
        'diverse, document, domain, domestic, dominate, draft, drama, duration, dynamic, economy',
        'edit, element, eliminate, emerge, emphasis, empirical, enable, encounter, energy, enforce',
        'enhance, enormous, ensure, entity, environment, equate, equip, equivalent, erode, error',
        'establish, estate, estimate, ethic, ethnic, evaluate, eventual, evident, evolve, exceed',
        'exclude, exclusive, exhibit, expand, expert, explicit, exploit, export, expose, external',
    ],
    'Day 11-13: 环境类词汇(1-90)': [
        'biodiversity, ecosystem, habitat, conservation, deforestation, degradation, emission, pollutant, carbon, greenhouse',
        'renewable, sustainable, fossil fuel, climate change, ozone layer, endangered, extinction, wetland, desertification, drought',
        'erosion, fertilizer, pesticide, herbicide, irrigation, organic, contamination, toxic, waste, recycle',
        'atmosphere, biosphere, ecology, precipitation, evaporation, condensation, runoff, groundwater, aquifer, watershed',
        'nitrogen cycle, carbon cycle, phosphorus, biomass, decomposition, photosynthesis, respiration, trophic, predator, prey',
        'mutualism, parasitism, commensalism, symbiosis, adaptation, natural selection, speciation, gene pool, allele, genotype',
        'phenotype, mutation, variation, inheritance, dominant, recessive, homozygous, heterozygous, polygenic, epigenetic',
        'conservation biology, restoration ecology, invasive species, endemic, flagship species, keystone species, corridor, fragmentation, reserve, buffer zone',
        'environmental impact, ecological footprint, carrying capacity, limiting factor, succession, climax community, pioneer species, biome, tundra, taiga',
    ],
    'Day 14-16: 科技与健康类词汇(1-90)': [
        'innovation, technology, automation, artificial intelligence, algorithm, database, network, interface, software, hardware',
        'digital, virtual, cyber, bandwidth, protocol, encryption, biotechnology, nanotechnology, renewable energy, solar panel',
        'wind turbine, geothermal, hydroelectric, nuclear, genome, clone, stem cell, genetic engineering, CRISPR, recombinant',
        'diagnosis, symptom, therapy, medication, vaccine, immunity, pathogen, virus, bacteria, infection',
        'chronic, acute, malignant, benign, metabolism, hormone, enzyme, receptor, antibody, antigen',
        'epidemiology, prevalence, incidence, mortality, morbidity, placebo, clinical trial, randomized, double-blind, cohort',
        'cardiovascular, respiratory, neurological, gastrointestinal, endocrine, immune system, inflammation, allergy, autoimmune, deficiency',
        'nutrient, vitamin, mineral, protein, carbohydrate, lipid, fiber, antioxidant, calorie, diet',
        'pharmaceutical, antibiotic, antiviral, analgesic, anesthetic, sedative, stimulant, antidepressant, anti-inflammatory, antihistamine',
    ],
    'Day 17-19: 健康类词汇续(91-180)': [
        'anatomy, physiology, pathology, pharmacology, toxicology, oncology, cardiology, neurology, dermatology, ophthalmology',
        'surgery, transplant, prosthesis, rehabilitation, physiotherapy, psychotherapy, cognitive, behavioral, developmental, genetic',
        'disorder, syndrome, condition, prognosis, remission, relapse, complication, contraindication, side effect, adverse reaction',
        'screening, prevention, intervention, treatment, palliative, curative, holistic, alternative medicine, acupuncture, herbal',
        'blood pressure, heart rate, respiration rate, body temperature, BMI, cholesterol, glucose, hemoglobin, platelet, white blood cell',
        'MRI, CT scan, ultrasound, X-ray, PET scan, biopsy, endoscopy, catheterization, dialysis, ventilation',
        'public health, sanitation, hygiene, epidemic, pandemic, quarantine, isolation, contact tracing, herd immunity, vaccination',
        'mental health, depression, anxiety, schizophrenia, bipolar disorder, PTSD, OCD, ADHD, autism, dementia',
        'lifestyle, exercise, meditation, stress management, sleep hygiene, addiction, substance abuse, alcoholism, smoking cessation, obesity',
    ],
    'Day 20-22: 教育类词汇(1-90)': [
        'curriculum, syllabus, pedagogy, didactic, heuristic, empirical, theoretical, practical, vocational, academic',
        'undergraduate, postgraduate, doctoral, dissertation, thesis, seminar, tutorial, lecture, workshop, laboratory',
        'scholarship, fellowship, grant, tuition, enrollment, admission, prerequisite, elective, compulsory, modular',
        'assessment, evaluation, examination, coursework, assignment, plagiarism, citation, reference, bibliography, peer review',
        'literacy, numeracy, critical thinking, problem-solving, creativity, collaboration, communication, autonomy, motivation, engagement',
        'cognitive, metacognitive, affective, psychomotor, Bloom taxonomy, scaffolding, differentiation, inclusion, diversity, equity',
        'primary education, secondary education, tertiary education, lifelong learning, distance learning, e-learning, blended learning, flipped classroom, MOOC, micro-credential',
        'accreditation, qualification, certification, diploma, degree, bachelor, master, doctor, honorary, professional',
        'discipline, interdisciplinary, multidisciplinary, transdisciplinary, specialization, major, minor, concentration, elective, core',
    ],
    'Day 23-25: 写作高频词汇(1-90)': [
        'moreover, furthermore, additionally, in addition, similarly, likewise, equally, correspondingly, analogously, comparably',
        'however, nevertheless, nonetheless, conversely, in contrast, on the contrary, whereas, while, although, despite',
        'therefore, consequently, accordingly, thus, hence, as a result, thereby, subsequently, inevitably, ultimately',
        'specifically, particularly, notably, especially, namely, that is, in other words, to clarify, to illustrate, for instance',
        'significantly, substantially, considerably, dramatically, markedly, marginally, minimally, moderately, exponentially, proportionally',
        'argue, contend, assert, maintain, claim, suggest, propose, advocate, emphasize, highlight',
        'demonstrate, illustrate, exemplify, elucidate, clarify, explicate, delineate, depict, portray, characterize',
        'fundamental, essential, crucial, vital, critical, pivotal, paramount, indispensable, integral, inherent',
        'potential, prospective, forthcoming, impending, imminent, eventual, ultimate, long-term, short-term, intermediate',
    ],
    'Day 26-28: 口语高频词汇(1-90)': [
        'absolutely, definitely, certainly, surely, undoubtedly, presumably, supposedly, apparently, seemingly, arguably',
        'basically, essentially, fundamentally, primarily, predominantly, principally, generally, typically, normally, usually',
        'fascinating, intriguing, captivating, compelling, engaging, appealing, enticing, alluring, charming, enchanting',
        'overwhelming, daunting, intimidating, formidable, challenging, demanding, strenuous, arduous, rigorous, grueling',
        'convenient, accessible, available, handy, practical, functional, efficient, effective, productive, streamlined',
        'nostalgic, reminiscent, evocative, sentimental, poignant, touching, moving, heartfelt, sincere, genuine',
        'trendy, fashionable, stylish, chic, elegant, sophisticated, refined, polished, sleek, contemporary',
        'vibrant, dynamic, lively, bustling, thriving, flourishing, prosperous, booming, burgeoning, emerging',
        'breathtaking, stunning, spectacular, magnificent, glorious, splendid, sublime, exquisite, picturesque, scenic',
    ],
    'Day 29-30: 真题高频词汇(1-60)': [
        'predominantly, invariably, fundamentally, intrinsically, inherently, ostensibly, purportedly, allegedly, notionally, hypothetically',
        'consolidate, reinforce, strengthen, bolster, enhance, augment, amplify, intensify, heighten, exacerbate',
        'mitigate, alleviate, ameliorate, relieve, ease, lessen, diminish, reduce, curb, suppress',
        'facilitate, expedite, accelerate, hasten, precipitate, catalyze, stimulate, promote, foster, cultivate',
        'scrutinize, examine, inspect, investigate, probe, explore, analyze, dissect, evaluate, assess',
        'synthesize, integrate, consolidate, unify, merge, amalgamate, coalesce, converge, reconcile, harmonize',
    ],
}

for title, words in ielts_vocab.items():
    add_heading_styled(title, 2)
    for line in words:
        add_bullet(line)

doc.add_page_break()

# ===== 生物专业词汇 =====
add_heading_styled('四、生物专业词汇表（450词）', 1)

bio_vocab = {
    'Cell & Molecular Biology': [
        'nucleus, mitochondria, ribosome, endoplasmic reticulum, Golgi apparatus, lysosome, peroxisome, cytoplasm, membrane, chloroplast',
        'chromosome, chromatin, nucleosome, histone, centromere, telomere, kinetochore, spindle, microtubule, microfilament',
        'transcription, translation, replication, DNA polymerase, RNA polymerase, helicase, ligase, primase, topoisomerase, exonuclease',
        'promoter, enhancer, silencer, operator, operon, repressor, activator, transcription factor, TATA box, CpG island',
        'mRNA, tRNA, rRNA, siRNA, miRNA, lncRNA, circRNA, snRNA, snoRNA, piRNA',
        'codon, anticodon, exon, intron, splicing, alternative splicing, polyadenylation, capping, ribozyme, telomerase',
        'phosphorylation, ubiquitination, acetylation, methylation, glycosylation, lipidation, SUMOylation, proteolysis, chaperone, proteasome',
        'signal transduction, receptor, ligand, kinase, phosphatase, second messenger, cAMP, calcium, G-protein, tyrosine kinase',
        'apoptosis, necrosis, autophagy, senescence, cell cycle, mitosis, meiosis, cytokinesis, checkpoint, cyclin',
    ],
    'Genetics & Genomics': [
        'allele, locus, genotype, phenotype, homozygous, heterozygous, dominant, recessive, codominant, epistasis',
        'segregation, independent assortment, linkage, recombination, crossing over, genetic map, physical map, QTL, GWAS, heritability',
        'mutation, point mutation, frameshift, deletion, insertion, duplication, inversion, translocation, aneuploidy, polyploidy',
        'genome, transcriptome, proteome, metabolome, epigenome, microbiome, metagenomics, sequencing, assembly, annotation',
        'SNP, indel, CNV, structural variant, haplotype, linkage disequilibrium, imputation, phasing, genotyping, phenotyping',
        'CRISPR, Cas9, gRNA, PAM, NHEJ, HDR, base editing, prime editing, knock-out, knock-in',
        'forward genetics, reverse genetics, mutagenesis, T-DNA, transposon, RNAi, gene silencing, overexpression, complementation, reporter gene',
        'pedigree, proband, consanguinity, autosomal, X-linked, Y-linked, mitochondrial, multifactorial, penetrance, expressivity',
        'population genetics, Hardy-Weinberg, genetic drift, gene flow, selection, fitness, bottleneck, founder effect, inbreeding, outbreeding',
    ],
    'Plant Breeding & Agriculture': [
        'germplasm, accession, landrace, cultivar, hybrid, inbred line, open-pollinated, synthetic, composite, clone',
        'heterosis, combining ability, GCA, SCA, diallel, topcross, testcross, backcross, recurrent selection, mass selection',
        'male sterility, CMS, GMS, restorer, maintainer, fertility, pollination, selfing, outcrossing, emasculation',
        'phenotyping, high-throughput, remote sensing, UAV, multispectral, hyperspectral, thermal, LiDAR, NDVI, canopy',
        'genomic selection, prediction model, training population, breeding value, GEBV, accuracy, cross-validation, BLUP, GBLUP, Bayesian',
        'yield, biomass, harvest index, thousand-grain weight, grain quality, protein content, starch, oil content, lodging, stress tolerance',
        'drought, salinity, heat, cold, frost, waterlogging, nutrient deficiency, disease resistance, pest resistance, herbicide tolerance',
        'marker-assisted selection, MABC, MARS, gene pyramiding, foreground selection, background selection, recombinant selection, flanking marker, diagnostic marker, KASP',
        'speed breeding, doubled haploid, embryo rescue, wide hybridization, somatic hybridization, protoplast fusion, cybrid, synthetic biology, gene stacking, seed production',
    ],
    'Biochemistry & Metabolism': [
        'glycolysis, Krebs cycle, oxidative phosphorylation, electron transport chain, ATP synthase, fermentation, gluconeogenesis, glycogen, pentose phosphate, Calvin cycle',
        'photosystem I, photosystem II, light reaction, dark reaction, RuBisCO, photorespiration, C3, C4, CAM, chlorophyll',
        'fatty acid, triglyceride, phospholipid, cholesterol, steroid, lipoprotein, beta-oxidation, ketogenesis, lipogenesis, desaturase',
        'amino acid, peptide, polypeptide, protein folding, alpha-helix, beta-sheet, domain, motif, disulfide bond, denaturation',
        'enzyme kinetics, Michaelis-Menten, Vmax, Km, competitive inhibition, noncompetitive, allosteric, cofactor, coenzyme, prosthetic group',
    ],
    'Microbiology & Pathology': [
        'Gram-positive, Gram-negative, peptidoglycan, lipopolysaccharide, flagellum, pilus, capsule, endospore, biofilm, quorum sensing',
        'pathogenicity, virulence, toxin, endotoxin, exotoxin, adhesion, invasion, colonization, host, vector',
        'fungus, hyphae, mycelium, spore, conidia, yeast, mold, dimorphic, saprophyte, symbiont',
        'virus, capsid, envelope, spike protein, reverse transcriptase, integrase, protease, retrovirus, bacteriophage, viroid',
        'immune response, innate immunity, adaptive immunity, T-cell, B-cell, macrophage, dendritic cell, cytokine, chemokine, interferon',
    ],
    'Bioinformatics & Statistics': [
        'algorithm, BLAST, alignment, FASTA, FASTQ, SAM, BAM, VCF, GFF, BED',
        'assembly, contig, scaffold, N50, coverage, depth, read, paired-end, mate-pair, long-read',
        'phylogenetic tree, neighbor-joining, maximum likelihood, bootstrap, clade, outgroup, ortholog, paralog, homolog, synteny',
        'p-value, confidence interval, standard deviation, standard error, t-test, ANOVA, chi-square, regression, correlation, PCA',
        'machine learning, random forest, SVM, neural network, deep learning, cross-validation, overfitting, regularization, feature selection, dimensionality reduction',
        'R package, Bioconductor, Python, pandas, NumPy, SciPy, scikit-learn, matplotlib, ggplot2, tidyverse',
    ],
}

for title, words in bio_vocab.items():
    add_heading_styled(title, 2)
    for line in words:
        add_bullet(line)

doc.add_page_break()

# ===== 语法点详解 =====
add_heading_styled('五、30天语法学习大纲', 1)

grammar_points = [
    ('Day 1', '一般现在时 vs 现在进行时', 'Simple Present (facts/habits) vs Present Continuous (actions now/temporary). Key: state verbs (know, believe, belong) rarely use continuous.'),
    ('Day 2', '一般过去时 vs 过去进行时', 'Simple Past (completed actions) vs Past Continuous (actions in progress at a past time). Often used together: "I was reading when she called."'),
    ('Day 3', '现在完成时 vs 过去完成时', 'Present Perfect (past with present relevance) vs Past Perfect (action before another past action). Signal words: already, yet, just, since, for.'),
    ('Day 4', '被动语态 - 一般时', 'am/is/are + past participle (present); was/were + past participle (past). Common in scientific writing: "The sample was heated."'),
    ('Day 5', '被动语态 - 完成时与进行时', 'has/have been + p.p.; had been + p.p.; is/are being + p.p. Example: "The experiment has been conducted."'),
    ('Day 6', '情态动词 (can/could/may/might)', 'Ability, permission, possibility. Can (present ability), could (past ability/polite), may (formal permission/possibility), might (weaker possibility).'),
    ('Day 7', '情态动词 (must/should/will/would)', 'Must (obligation/strong deduction), should (advice), will (future/prediction), would (hypothetical/polite).'),
    ('Day 8', '条件句 - 零条件与第一条件', 'Zero: If + present, present (facts). First: If + present, will + base (real future). "If it rains, the soil will be saturated."'),
    ('Day 9', '条件句 - 第二与第三条件', 'Second: If + past, would + base (unreal present). Third: If + past perfect, would have + p.p. (unreal past). Keys for hypothesis in Discussion sections.'),
    ('Day 10', '关系从句 (who/which/that)', 'Defining vs non-defining relative clauses. Who (people), which (things), that (both, defining only). "The gene that controls flowering..."'),
    ('Day 11', '关系从句 (where/when/whose)', 'Where (place), when (time), whose (possession). "The lab where we conducted PCR..."'),
    ('Day 12', '不定式与动名词', 'to + verb vs verb-ing. After certain verbs: enjoy doing, want to do. After prepositions: interested in doing. Purpose: to + verb.'),
    ('Day 13', '冠词 (a/an/the) 精讲', 'Indefinite (a/an: first mention, general) vs definite (the: specific, unique, second mention). Zero article: plural/general, abstract nouns, proper nouns.'),
    ('Day 14', '可数/不可数名词', 'Countable: singular/plural, a/an. Uncountable: no a/an, no plural (information, equipment, research, data). Quantifiers: many/much, few/little.'),
    ('Day 15', '比较级与最高级', '-er/more + than; the -est/the most. Irregular: good-better-best, bad-worse-worst. As...as. "The yield was significantly higher than the control."'),
    ('Day 16', '间接引语 - 陈述句', 'Tense backshift: present→past, past→past perfect. Pronoun changes. "The researcher said (that) the results were significant."'),
    ('Day 17', '间接引语 - 疑问句与祈使句', 'Yes/No questions: if/whether. Wh-questions: keep question word. Word order becomes statement. "She asked whether the PCR had worked."'),
    ('Day 18', '虚拟语气 (wish/if only)', 'Wish + past (unreal present): "I wish I had more data." Wish + past perfect (regret): "I wish I had used a larger sample size."'),
    ('Day 19', '倒装句', 'Negative adverbs fronting: Never/Not only/Seldom + auxiliary + subject + verb. "Not only does temperature affect growth, but it also..."'),
    ('Day 20', '强调句 (It is...that/cleft sentences)', 'It is/was + emphasized element + that/who... "It is the interaction effect that drives the phenotype." What-clause: "What we found was..."'),
    ('Day 21', '并列连词与从属连词', 'Coordinating: and, but, or, nor, for, so, yet. Subordinating: because, although, while, whereas, since, unless, until, after, before.'),
    ('Day 22', '介词搭配 - 时间与地点', 'At (point), in (enclosed), on (surface). At 3pm, in July, on Monday. At the lab, in the field, on the plate.'),
    ('Day 23', '介词搭配 - 动词+介词', 'Depend on, result in, contribute to, consist of, associate with, differ from, compare to/with, relate to, focus on, lead to.'),
    ('Day 24', '主谓一致', 'Singular/plural agreement. Collective nouns. Either/or, neither/nor. "The data are/were..." (plural, though data is is also accepted).'),
    ('Day 25', '定语从句 vs 分词短语', 'Reducing relative clauses: "The gene encoding the protein" (= that encodes). "The plants grown in the greenhouse" (= that were grown).'),
    ('Day 26', '名词性从句', 'That-clause: "We hypothesize that..." Wh-clause: "We investigated whether..." Subject clause: "What remains unclear is..."'),
    ('Day 27', '状语从句全总结', 'Time (when, while, as soon as), reason (because, since, as), purpose (so that, in order that), concession (although, even though), result (so...that).'),
    ('Day 28', '学术写作常用句型 (1)', 'Describing methods: "X was performed/carried out/conducted using..." Presenting results: "X increased/decreased significantly (p<0.05)." "As shown in Figure 1..."'),
    ('Day 29', '学术写作常用句型 (2)', 'Discussion: "These findings suggest/indicate/demonstrate that..." "This is consistent with previous studies..." "One limitation of this study is..."'),
    ('Day 30', '语法总复习 + 常见错误', 'Top errors: subject-verb agreement, article usage, tense consistency, preposition choice, comma splices, dangling modifiers. Self-check checklist.'),
]

for day, topic, detail in grammar_points:
    add_heading_styled(f'{day}: {topic}', 2)
    add_para(detail)

doc.add_page_break()

# ===== 生物实验英文流程 =====
add_heading_styled('六、英文生物实验标准流程（每3天1个，共10个）', 1)

protocols = [
    ('Protocol 1: Genomic DNA Extraction from Plant Tissue (CTAB Method)', [
        '1. Grind 100 mg of fresh leaf tissue in liquid nitrogen to a fine powder.',
        '2. Transfer powder to a 2 mL microcentrifuge tube containing 700 μL preheated (65°C) CTAB extraction buffer.',
        '3. Add 2 μL of β-mercaptoethanol and vortex thoroughly. Incubate at 65°C for 45-60 min, inverting every 10 min.',
        '4. Add 700 μL of chloroform:isoamyl alcohol (24:1), vortex for 30 sec, and centrifuge at 12,000 rpm for 10 min at room temperature.',
        '5. Transfer the aqueous (upper) phase to a new 1.5 mL tube. Add 0.7 volumes of cold isopropanol, mix gently by inversion.',
        '6. Incubate at -20°C for 30 min, then centrifuge at 12,000 rpm for 10 min at 4°C.',
        '7. Discard supernatant. Wash pellet with 500 μL of 70% ethanol, centrifuge at 12,000 rpm for 5 min.',
        '8. Repeat wash once. Air-dry pellet for 10-15 min (do not over-dry).',
        '9. Resuspend DNA in 50-100 μL TE buffer or nuclease-free water. Add 1 μL RNase A (10 mg/mL), incubate at 37°C for 30 min.',
        '10. Quantify DNA concentration using NanoDrop spectrophotometer (A260/A280 ratio 1.8-2.0). Store at -20°C.',
        '', 'Safety Notes: Work in a fume hood when handling chloroform and β-mercaptoethanol. Wear gloves and lab coat.',
    ]),
    ('Protocol 2: Polymerase Chain Reaction (PCR)', [
        '1. Prepare PCR master mix on ice in a 0.2 mL PCR tube (total volume 25-50 μL): 1× PCR buffer, 1.5-2.5 mM MgCl₂, 200 μM each dNTP, 0.2-0.5 μM each primer (forward and reverse), 1-2 U Taq DNA polymerase, 10-100 ng template DNA, nuclease-free water to final volume.',
        '2. Mix gently by pipetting or brief vortex. Spin down briefly to collect liquid.',
        '3. Place tube in thermocycler and run program: Initial denaturation: 95°C for 3-5 min; 30-35 cycles of: Denaturation: 95°C for 30 sec, Annealing: 50-65°C (primer Tm-dependent) for 30 sec, Extension: 72°C for 30-60 sec/kb; Final extension: 72°C for 5-10 min; Hold at 4°C.',
        '4. Verify amplification by agarose gel electrophoresis (1-2% gel, stain with ethidium bromide or SYBR Safe).',
        '5. Include appropriate controls: positive control (known template), negative control (no template / water blank).',
        '', 'Troubleshooting: No bands → check primer design, increase Mg²⁺, reduce annealing temperature. Non-specific bands → increase annealing temperature, reduce Mg²⁺, reduce cycles.',
    ]),
    ('Protocol 3: Agarose Gel Electrophoresis', [
        '1. Prepare 1× TAE or TBE running buffer. Determine gel percentage based on expected DNA size (0.8% for large fragments >2 kb; 1-2% for 0.5-2 kb; 2-3% for <500 bp).',
        '2. Weigh appropriate amount of agarose powder and add to 1× buffer in a flask. Heat in microwave until completely dissolved (clear solution, no particles).',
        '3. Cool solution to ~55°C (can hold flask under running water). Add DNA stain (e.g., ethidium bromide 0.5 μg/mL or SYBR Safe 1×). Swirl gently to mix.',
        '4. Pour gel into casting tray with comb in place. Allow to solidify at room temperature for 20-30 min (gel should appear opaque).',
        '5. Place gel in electrophoresis tank, cover with 1× running buffer. Carefully remove comb.',
        '6. Mix DNA samples with 6× loading dye (final 1×). Load samples and DNA ladder (e.g., 1 kb or 100 bp ladder) into wells.',
        '7. Run at 5-10 V/cm (distance between electrodes) until dye front migrates 2/3 to 3/4 of gel length (typically 30-60 min).',
        '8. Visualize under UV transilluminator or blue light transilluminator. Document with gel imaging system.',
        '', 'Safety: Ethidium bromide is a mutagen. Use nitrile gloves. Dispose of contaminated materials in designated waste.',
    ]),
    ('Protocol 4: Bacterial Transformation (Heat Shock Method)', [
        '1. Thaw competent E. coli cells (e.g., DH5α) on ice for 10-15 min (do not vortex or warm with hands).',
        '2. Add 1-5 μL of ligation product or 1-10 ng of plasmid DNA to 50 μL of competent cells. Mix gently by flicking tube (do not pipette up and down vigorously).',
        '3. Incubate on ice for 30 min.',
        '4. Heat shock at exactly 42°C for 45-90 seconds (optimize based on tube type and cell batch). Place immediately back on ice for 2 min.',
        '5. Add 450-950 μL of pre-warmed (37°C) SOC or LB medium (without antibiotic).',
        '6. Incubate at 37°C with shaking (200-250 rpm) for 45-60 min (recovery/expression period).',
        '7. Plate 50-200 μL of transformation mix on pre-warmed LB agar plates containing appropriate antibiotic(s). Spread evenly with sterile glass spreader or beads.',
        '8. Centrifuge remaining culture at 4,000 rpm for 2 min, resuspend in ~100 μL supernatant, and plate as a concentrated sample.',
        '9. Incubate plates inverted at 37°C overnight (14-16 h).',
        '10. Check for colonies the next morning. Pick single colonies for further analysis (colony PCR, plasmid extraction, sequencing).',
        '', 'Controls: Positive control (known plasmid), negative control (no DNA / TE only). Key: Pre-warm plates and SOC medium to 37°C before use.',
    ]),
    ('Protocol 5: Total RNA Extraction (TRIzol Method)', [
        '1. Homogenize 50-100 mg of tissue in 1 mL TRIzol reagent using mortar and pestle (liquid nitrogen) or homogenizer. Work quickly on ice.',
        '2. Incubate homogenate at room temperature for 5 min to permit complete dissociation of nucleoprotein complexes.',
        '3. Add 200 μL of chloroform per 1 mL TRIzol. Shake tube vigorously by hand for 15 sec. Incubate at room temperature for 2-3 min.',
        '4. Centrifuge at 12,000 × g for 15 min at 4°C. The mixture separates into 3 phases: lower red phenol-chloroform phase (proteins), white interphase (DNA), upper colorless aqueous phase (RNA).',
        '5. Carefully transfer ~400-500 μL of aqueous phase to a new RNase-free tube (avoid touching interphase).',
        '6. Add 500 μL of isopropanol per 1 mL TRIzol used initially. Mix by inversion. Incubate at room temperature for 10 min.',
        '7. Centrifuge at 12,000 × g for 10 min at 4°C. RNA pellet should be visible as a white/gel-like pellet.',
        '8. Remove supernatant. Wash pellet with 1 mL of 75% ethanol (prepared with DEPC-treated water). Centrifuge at 7,500 × g for 5 min at 4°C.',
        '9. Discard supernatant. Air-dry pellet for 5-10 min (do not over-dry; pellet becomes transparent).',
        '10. Resuspend in 30-50 μL DEPC-treated water or TE buffer. Heat at 55-60°C for 10-15 min if needed for dissolution.',
        '11. Quantify RNA (A260/280 ratio ~2.0; A260/230 ratio >1.8). Check integrity on 1% denaturing agarose gel or Bioanalyzer. Store at -80°C.',
        '', 'Critical: Use RNase-free consumables and DEPC-treated water throughout. Work in RNA-dedicated area. Change gloves frequently.',
    ]),
    ('Protocol 6: Quantitative Real-Time PCR (qPCR)', [
        '1. Design primers using software (Primer3, NCBI Primer-BLAST): amplicon 80-150 bp, Tm ~60°C, GC content 40-60%, avoid secondary structures. Validate primer specificity by BLAST and melt curve analysis.',
        '2. Synthesize cDNA from 1 μg total RNA using reverse transcriptase kit (random hexamer or oligo-dT priming) following manufacturer protocol. Dilute cDNA 1:5 to 1:10.',
        '3. Prepare qPCR reaction mix per well (10-20 μL total): 1× SYBR Green Master Mix, 200-400 nM each primer, 1-2 μL diluted cDNA, nuclease-free water to volume. Include reference dye (ROX) if required by instrument.',
        '4. Plate layout: each sample in triplicate, include no-template controls (NTC), and reference gene(s) for normalization. Use optical adhesive seal.',
        '5. Centrifuge plate at 1,000 rpm for 1 min to remove bubbles and collect liquid.',
        '6. Run program: 95°C for 2-10 min (polymerase activation); 40 cycles of: 95°C for 15 sec, 60°C for 30-60 sec (combined annealing/extension); Melt curve: 65-95°C in 0.5°C increments.',
        '7. Analyze data using the 2^(-ΔΔCt) method: ΔCt = Ct(target) - Ct(reference); ΔΔCt = ΔCt(treatment) - ΔCt(control); Fold change = 2^(-ΔΔCt).',
        '8. Verify specificity by single-peak melt curves. Check NTC for contamination. Report Ct SD <0.5 among technical replicates.',
        '', 'Reference Gene Selection: Test at least 3 candidate reference genes (e.g., ACTIN, GAPDH, 18S rRNA, UBQ, EF1α) for expression stability across your experimental conditions.',
    ]),
    ('Protocol 7: SDS-PAGE and Western Blotting', [
        '1. Prepare protein samples: extract total protein using RIPA buffer with protease/phosphatase inhibitors. Quantify by BCA or Bradford assay. Add 4× Laemmli loading buffer, boil at 95°C for 5 min.',
        '2. Prepare SDS-PAGE gel: pour resolving gel (e.g., 10-12% acrylamide), overlay with water/isopropanol, let polymerize 30 min. Pour stacking gel (5%), insert comb, polymerize 30 min.',
        '3. Load 20-50 μg protein per lane plus pre-stained protein ladder. Run at 80 V through stacking gel, then 120 V through resolving gel until dye front reaches bottom (~1.5 h).',
        '4. Transfer: equilibrate gel in transfer buffer 10 min. Assemble sandwich: sponge → filter paper → gel → PVDF (pre-activated in methanol 1 min) or nitrocellulose membrane → filter paper → sponge. Remove all bubbles.',
        '5. Transfer at 100 V for 1-1.5 h at 4°C (or 30 V overnight). Verify transfer by Ponceau S staining or pre-stained ladder on membrane.',
        '6. Block membrane in 5% non-fat dry milk or BSA in TBST for 1 h at room temperature with gentle shaking.',
        '7. Incubate with primary antibody (diluted in blocking buffer, typically 1:500-1:5000) overnight at 4°C with gentle shaking.',
        '8. Wash 3× 10 min with TBST. Incubate with HRP-conjugated secondary antibody (1:5000-1:20000) for 1 h at room temperature.',
        '9. Wash 3× 10 min with TBST. Detect signal using ECL substrate. Image with chemiluminescence imaging system.',
        '10. Strip and re-probe for loading control (e.g., β-actin, GAPDH) if needed. Quantify band intensity using ImageJ.',
    ]),
    ('Protocol 8: Agrobacterium-Mediated Plant Transformation (Leaf Disc Method)', [
        '1. Prepare Agrobacterium tumefaciens culture: inoculate single colony in 5 mL LB with appropriate antibiotics. Grow overnight at 28°C, 200 rpm.',
        '2. Subculture 1:50 in fresh LB + antibiotics + 100 μM acetosyringone. Grow to OD600 = 0.5-0.8. Centrifuge at 4,000 rpm for 10 min. Resuspend in MS liquid medium + 100 μM acetosyringone to OD600 = 0.5.',
        '3. Prepare explants: cut young leaves from 4-6 week old sterile-grown plants into 0.5-1 cm² discs. Pre-culture on MS + 2 mg/L 6-BA + 0.1 mg/L NAA for 2 days in dark.',
        '4. Inoculate: immerse leaf discs in Agrobacterium suspension for 10-15 min with gentle shaking. Blot excess on sterile filter paper.',
        '5. Co-cultivation: place discs (adaxial side up) on co-cultivation medium (MS + hormones + 100 μM acetosyringone). Incubate in dark at 25°C for 2-3 days.',
        '6. Wash discs 3-5 times with sterile water containing 500 mg/L cefotaxime (to kill Agrobacterium). Blot dry.',
        '7. Selection: transfer to shoot regeneration medium (MS + hormones + selection antibiotic + 500 mg/L cefotaxime). Culture under 16/8 h light/dark at 25°C. Subculture every 2 weeks.',
        '8. After 4-8 weeks, excise regenerated shoots (~1-2 cm) and transfer to rooting medium (1/2 MS + 0.1 mg/L IBA + selection antibiotic + cefotaxime).',
        '9. After roots develop (2-4 weeks), transfer plantlets to soil. Cover with plastic dome for 1 week for acclimatization. Gradually reduce humidity.',
        '10. Confirm transformation by: (a) reporter gene assay (GUS staining/GFP fluorescence), (b) PCR for transgene, (c) Southern blot for copy number, (d) qRT-PCR for expression level.',
    ]),
    ('Protocol 9: Enzyme-Linked Immunosorbent Assay (ELISA)', [
        '1. Coating: dilute capture antibody in coating buffer (0.05 M carbonate-bicarbonate, pH 9.6) to 1-10 μg/mL. Add 100 μL per well to 96-well plate. Seal and incubate overnight at 4°C.',
        '2. Wash 3× with 300 μL PBST (PBS + 0.05% Tween-20) per well. After each wash, tap plate firmly on paper towels to remove residual buffer.',
        '3. Blocking: add 200-300 μL blocking buffer (1-5% BSA or non-fat milk in PBS) per well. Incubate at room temperature for 1-2 h or overnight at 4°C.',
        '4. Wash 3× with PBST.',
        '5. Add 100 μL of standards and samples (diluted in blocking buffer) to appropriate wells. Include blank wells (blocking buffer only). Run each in duplicate. Incubate at room temperature for 2 h or overnight at 4°C.',
        '6. Wash 3-5× with PBST (thorough washing is critical for low background).',
        '7. Add 100 μL detection antibody (diluted in blocking buffer, typically 0.5-2 μg/mL) to each well. Incubate at room temperature for 1-2 h.',
        '8. Wash 3-5× with PBST.',
        '9. Add 100 μL enzyme conjugate (e.g., streptavidin-HRP) diluted in blocking buffer. Incubate at room temperature for 30 min-1 h in dark.',
        '10. Wash 5-7× with PBST.',
        '11. Add 100 μL TMB substrate solution. Incubate in dark at room temperature for 15-30 min (monitor color development; blue color should appear in positive wells).',
        '12. Stop reaction with 50-100 μL stop solution (2 M H₂SO₄ or 1 M HCl). Color turns yellow. Read absorbance at 450 nm within 30 min. Use 570 nm as reference wavelength.',
        '13. Generate standard curve (4-parameter logistic fit). Calculate sample concentrations from standard curve. Ensure R² > 0.99 for standard curve.',
    ]),
    ('Protocol 10: Field Trial Design for Crop Breeding (RCBD)', [
        '1. Define experimental objective and treatments. Select Randomized Complete Block Design (RCBD) as default for field heterogeneity. Determine number of replications (minimum 3, ideally 4-6).',
        '2. Site preparation: soil testing (pH, N-P-K, organic matter), uniform tillage, pre-planting irrigation if needed. Map field for slope, drainage, and known variability.',
        '3. Layout: divide field into blocks perpendicular to the gradient of variability. Each block contains one plot of each genotype/treatment in random order. Plot size: typically 4-6 rows × 5 m length for cereals. Leave alleyways (0.5-1 m) between blocks.',
        '4. Randomization: use random number generator or field design software (e.g., agricolae package in R, DiGGer). Assign treatment codes to each plot within each block independently.',
        '5. Planting: label each plot clearly with weatherproof tags. Record planting date, seeding rate, row spacing. Plant border rows around the entire experiment to minimize edge effects.',
        '6. Management: apply uniform fertilizer, irrigation, and pest/disease control across all plots unless testing management treatments. Record all inputs (type, rate, date).',
        '7. Data collection: measure traits at appropriate growth stages. Record: (a) emergence date/rate, (b) flowering date (50% anthesis), (c) plant height, (d) chlorophyll content (SPAD), (e) canopy temperature, (f) yield and yield components at harvest, (g) grain quality parameters.',
        '8. Harvest: harvest central rows only to avoid border effects (discard at least one border row per side). Record plot harvest area. Measure fresh weight, then dry to constant weight for dry matter yield.',
        '9. Statistical analysis: ANOVA with block as random effect and genotype as fixed effect. Test for normality (Shapiro-Wilk) and homogeneity of variance (Bartlett/Levene). Transform data if needed. Calculate broad-sense heritability: H² = σ²g / (σ²g + σ²e/r).',
        '10. Report: present ANOVA table with F-values and p-values. Use LSD or Tukey HSD for mean separation (α = 0.05). Include CV% to indicate experimental precision (<15% acceptable for yield trials).',
    ]),
]

for title, steps in protocols:
    add_heading_styled(title, 2)
    for step in steps:
        if step.startswith('Safety') or step.startswith('Critical') or step.startswith('Troubleshooting') or step.startswith('Reference') or step.startswith('Controls'):
            add_para(step, bold=True)
        elif step == '':
            doc.add_paragraph()
        else:
            add_para(step, indent=False)

doc.add_page_break()

# ===== 学习方法建议 =====
add_heading_styled('七、学习方法与建议', 1)

add_para('1. 词汇记忆法：采用"艾宾浩斯遗忘曲线"复习法——新词学习后在第1、2、4、7、15天重复复习。每天早上30分钟新词，晚上15分钟复习旧词。')
add_para('2. 词汇卡片：准备两种颜色卡片——黄色记雅思词汇，绿色记生物词汇。正面写英文，背面写中文释义+例句。随身携带，利用碎片时间。')
add_para('3. 语法学习：每天先阅读语法点讲解（15分钟），然后做10道相关练习题（15分钟）。每周日回顾本周所有语法点，整理错题。')
add_para('4. 实验流程：学习每个Protocol时，先朗读一遍熟悉发音，再默写关键步骤（如试剂名称和浓度），最后用自己的话口头复述整个流程。')
add_para('5. 真题应用：每周末做一套雅思阅读真题（学术类），注意标注文章中出现的本周词汇，加深记忆。')
add_para('6. 输出练习：每天用英语写一篇50-100字的简短日记，尝试使用当天学到的新词和语法点。')

# ===== 每日时间分配 =====
add_heading_styled('八、每日时间分配建议（约2.5小时）', 1)

make_table(
    ['时间段', '内容', '时长'],
    [
        ['早晨 (7:00-7:30)', '雅思词汇（30词）学习+朗读', '30分钟'],
        ['上午 (10:00-10:30)', '生物专业词汇（15词）学习', '30分钟'],
        ['下午 (14:00-14:30)', '语法点学习+练习', '30分钟'],
        ['晚间 (20:00-20:20)', '旧词复习（艾宾浩斯法）', '20分钟'],
        ['晚间 (20:20-20:50)', '每3天：英文实验流程学习（替换词汇复习）', '30分钟'],
        ['睡前 (22:00-22:10)', '手机APP刷词/听力泛听', '10分钟'],
    ]
)

doc.add_paragraph()

# ===== 推荐资源 =====
add_heading_styled('九、推荐学习资源', 1)

add_para('雅思备考：', bold=True, indent=False)
add_bullet('书籍：《剑桥雅思真题集》系列（Cambridge IELTS 10-18）')
add_bullet('APP：墨墨背单词、百词斩（雅思词库）')
add_bullet('网站：IELTS Liz (ieltsliz.com), IELTS Advantage (ieltsadvantage.com)')

add_para('生物专业英语：', bold=True, indent=False)
add_bullet('期刊：Nature, Science, Nature Genetics, Plant Cell（每周精读1篇Abstract+Introduction）')
add_bullet('教材：《Molecular Biology of the Gene》(Watson), 《Principles of Genetics》(Snustad)')
add_bullet('MOOC：Coursera - "Bioinformatics Specialization", edX - "Principles of Biochemistry"')

add_para('语法参考：', bold=True, indent=False)
add_bullet('《English Grammar in Use》(Raymond Murphy) — 中级（蓝皮）和高级（绿皮）两册')
add_bullet('《Academic Writing for Graduate Students》(Swales & Feak) — 学术写作圣经')

# 保存
output_path = r'D:\2026-SP\30天雅思与生物英语学习计划.docx'
doc.save(output_path)
print(f'✅ 输出: {output_path}')
print(f'   文件大小: {os.path.getsize(output_path)/1024:.0f} KB')
